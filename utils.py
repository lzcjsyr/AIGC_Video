"""
智能视频制作系统 - 工具函数模块
包含通用工具函数、错误处理和日志管理
"""

import os
import re
import json
import logging
import datetime
from pathlib import Path
from typing import Dict, Any, Optional, List

# 配置日志
logging.basicConfig(
    level=logging.INFO,
    format='%(asctime)s - %(name)s - %(levelname)s - %(message)s',
    handlers=[
        logging.FileHandler('aigc_video.log', encoding='utf-8'),
        logging.StreamHandler()
    ]
)

logger = logging.getLogger('AIGC_Video')

# 降低第三方库 pdfminer 的噪声日志级别
for _name in [
    "pdfminer",
    "pdfminer.pdffont",
    "pdfminer.pdfinterp",
    "pdfminer.cmapdb",
]:
    try:
        logging.getLogger(_name).setLevel(logging.ERROR)
    except Exception:
        pass

class VideoProcessingError(Exception):
    """视频处理专用异常类"""
    pass

class APIError(Exception):
    """API调用异常类"""
    pass

class FileProcessingError(Exception):
    """文件处理异常类"""
    pass

def log_function_call(func):
    """装饰器：记录函数调用"""
    def wrapper(*args, **kwargs):
        logger.info(f"调用函数: {func.__name__}")
        try:
            result = func(*args, **kwargs)
            logger.info(f"函数 {func.__name__} 执行成功")
            return result
        except Exception as e:
            logger.debug(f"函数 {func.__name__} 执行失败: {str(e)}")
            raise
    return wrapper

def ensure_directory_exists(directory: str) -> None:
    """确保目录存在，如不存在则创建"""
    Path(directory).mkdir(parents=True, exist_ok=True)
    logger.debug(f"确保目录存在: {directory}")

def clean_text(text: str) -> str:
    """清理文本内容"""
    if not text:
        return ""
    
    # 移除HTML标签
    text = re.sub(r'<[^>]+>', '', text)
    
    # 清理PDF CID字符乱码：移除 (cid:数字) 格式的字符
    text = re.sub(r'\(cid:\d+\)', '', text)
    
    # 清理其他常见的PDF解析问题
    # 移除单独的数字和字母组合（可能是字体编码残留）
    text = re.sub(r'\b[A-Z]{1,3}\d*\b', ' ', text)
    
    # 更强力的乱码字符清理
    # 移除明显的非文本字符（保留中文、英文、数字、常见标点）
    def is_valid_char(char):
        # 中文字符
        if '\u4e00' <= char <= '\u9fff':
            return True
        # 英文字母和数字
        if char.isalnum() and ord(char) < 128:
            return True
        # 常见标点符号
        if char in '，。！？；：""''（）【】《》、—…·.,:;!?()[]{}"-\'':
            return True
        # 空格和换行
        if char in ' \n\t\r':
            return True
        return False
    
    # 字符级过滤
    filtered_chars = []
    for char in text:
        if is_valid_char(char):
            filtered_chars.append(char)
        else:
            # 用空格替换无效字符
            if filtered_chars and filtered_chars[-1] != ' ':
                filtered_chars.append(' ')
    
    text = ''.join(filtered_chars)
    
    # 标准化空白字符
    text = re.sub(r'\s+', ' ', text)
    # 移除首尾空白
    text = text.strip()
    
    return text

def validate_file_format(file_path: str, supported_formats: List[str]) -> bool:
    """验证文件格式"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    file_extension = Path(file_path).suffix.lower()
    if file_extension not in supported_formats:
        raise FileProcessingError(f"不支持的文件格式: {file_extension}，支持的格式: {supported_formats}")
    
    return True

def safe_json_loads(json_string: str) -> Dict[str, Any]:
    """安全的JSON解析"""
    try:
        # 提取JSON部分
        json_start = json_string.find('{')
        json_end = json_string.rfind('}') + 1
        
        if json_start == -1 or json_end == 0:
            raise ValueError("未找到有效的JSON对象")
        
        json_content = json_string[json_start:json_end]
        return json.loads(json_content)
    
    except json.JSONDecodeError as e:
        logger.error(f"JSON解析失败: {str(e)}")
        raise ValueError(f"JSON格式错误: {str(e)}")

def parse_json_robust(raw_text: str) -> Dict[str, Any]:
    """鲁棒解析：先尝试标准JSON解析，失败再用json-repair做保守修复。
    - 仅对首次'{'与末次'}'之间的子串进行修复，避免引入额外内容
    - 修复成功后再用json.loads确认为有效JSON
    """
    # 提取JSON主体
    start = raw_text.find('{')
    end = raw_text.rfind('}')
    if start == -1 or end == -1 or end < start:
        raise ValueError("未在输出中找到 JSON 对象")
    snippet = raw_text[start:end+1]
    # 1) 常规解析
    try:
        return json.loads(snippet)
    except Exception as e1:
        logger.warning(f"标准JSON解析失败，尝试修复: {e1}")
    # 2) 尝试使用 json-repair 做保守修复
    try:
        from json_repair import repair_json
    except Exception as ie:
        raise ValueError(f"JSON解析失败，且缺少json-repair依赖: {ie}")
    try:
        repaired = repair_json(snippet, ensure_ascii=False)
        return json.loads(repaired)
    except Exception as e2:
        preview = snippet[:300]
        raise ValueError(f"JSON修复解析失败: {e2}; 片段预览: {preview}")

def save_json_file(data: Dict[str, Any], file_path: str) -> None:
    """安全地保存JSON文件"""
    try:
        ensure_directory_exists(os.path.dirname(file_path))
        with open(file_path, 'w', encoding='utf-8') as f:
            json.dump(data, f, ensure_ascii=False, indent=2)
        logger.info(f"JSON文件已保存: {file_path}")
    except Exception as e:
        logger.error(f"保存JSON文件失败 {file_path}: {str(e)}")
        raise FileProcessingError(f"保存文件失败: {str(e)}")

def load_json_file(file_path: str) -> Dict[str, Any]:
    """安全地加载JSON文件"""
    try:
        with open(file_path, 'r', encoding='utf-8') as f:
            data = json.load(f)
        logger.info(f"JSON文件已加载: {file_path}")
        return data
    except Exception as e:
        logger.error(f"加载JSON文件失败 {file_path}: {str(e)}")
        raise FileProcessingError(f"加载文件失败: {str(e)}")

def calculate_duration(text_length: int, speech_speed_wpm: int = 300) -> float:
    """计算文本播放时长（秒）"""
    # 中文按每分钟300字计算
    duration_seconds = (text_length / speech_speed_wpm) * 60
    return round(duration_seconds, 1)

def format_file_size(size_bytes: int) -> str:
    """格式化文件大小显示"""
    if size_bytes == 0:
        return "0B"
    
    size_names = ["B", "KB", "MB", "GB"]
    i = 0
    while size_bytes >= 1024 and i < len(size_names) - 1:
        size_bytes /= 1024.0
        i += 1
    
    return f"{size_bytes:.1f}{size_names[i]}"

def get_file_info(file_path: str) -> Dict[str, Any]:
    """获取文件信息"""
    if not os.path.exists(file_path):
        raise FileNotFoundError(f"文件不存在: {file_path}")
    
    stat = os.stat(file_path)
    
    return {
        "path": file_path,
        "name": os.path.basename(file_path),
        "size": stat.st_size,
        "size_formatted": format_file_size(stat.st_size),
        "modified_time": datetime.datetime.fromtimestamp(stat.st_mtime),
        "extension": Path(file_path).suffix.lower()
    }

def retry_on_failure(max_retries: int = 3, delay: float = 1.0):
    """重试装饰器"""
    def decorator(func):
        def wrapper(*args, **kwargs):
            last_exception = None
            
            for attempt in range(max_retries):
                try:
                    return func(*args, **kwargs)
                except Exception as e:
                    last_exception = e
                    if attempt < max_retries - 1:
                        logger.warning(f"函数 {func.__name__} 第{attempt + 1}次尝试失败: {str(e)}，{delay}秒后重试...")
                        import time
                        time.sleep(delay)
                    else:
                        logger.error(f"函数 {func.__name__} 经过{max_retries}次尝试后仍然失败")
            
            raise last_exception
        return wrapper
    return decorator

def validate_required_fields(data: Dict[str, Any], required_fields: List[str]) -> None:
    """验证必需字段"""
    missing_fields = []
    for field in required_fields:
        if field not in data or data[field] is None:
            missing_fields.append(field)
    
    if missing_fields:
        raise ValueError(f"缺少必需字段: {', '.join(missing_fields)}")

def create_processing_summary(
    input_file: str,
    original_length: int,
    target_length: int,
    actual_length: int,
    num_segments: int,
    start_time: datetime.datetime,
    end_time: datetime.datetime
) -> str:
    """创建处理摘要文本"""
    
    execution_time = (end_time - start_time).total_seconds()
    compression_ratio = (1 - actual_length / original_length) * 100 if original_length > 0 else 0
    
    summary = f"""=== 文档处理摘要 ===
处理时间: {start_time.strftime('%Y-%m-%d %H:%M:%S')}
原始文档: {os.path.basename(input_file)}
原始字数: {original_length:,}字
目标字数: {target_length}字
实际字数: {actual_length}字
压缩比例: {compression_ratio:.1f}%
分段数量: {num_segments}段
总处理时间: {execution_time:.1f}秒

=== 处理统计 ===
平均每段字数: {actual_length // num_segments}字
预估总播放时长: {calculate_duration(actual_length):.1f}秒
压缩效率: {((original_length - actual_length) / execution_time):.0f}字/秒
"""
    
    return summary

def progress_callback(current: int, total: int, operation: str = "处理"):
    """进度回调函数"""
    progress = (current / total) * 100 if total > 0 else 0
    logger.info(f"{operation}进度: {current}/{total} ({progress:.1f}%)")

def prompt_yes_no(message: str, default: bool = True) -> bool:
    """命令行确认提示，返回布尔。
    
    Args:
        message: 提示消息
        default: 默认选择（回车时采用）
    """
    try:
        suffix = "[Y/n]" if default else "[y/N]"
        # 统一在提示前输出一个空行，避免在调用点散落打印
        print()
        while True:
            choice = input(f"{message} {suffix}: ").strip().lower()
            if choice == '' and default is not None:
                return default
            if choice in ['y', 'yes', '是']:
                return True
            if choice in ['n', 'no', '否']:
                return False
            print("请输入 y 或 n")
    except KeyboardInterrupt:
        print("\n操作已取消")
        return False

def prompt_choice(message: str, options: List[str], default_index: int = 0) -> Optional[str]:
    """通用选项选择器，返回所选项文本。
    支持输入序号或精确匹配选项文本（不区分大小写）。
    """
    try:
        while True:
            print(f"\n{message}（输入 q 返回上一级）")
            for i, opt in enumerate(options, 1):
                prefix = "*" if (i - 1) == default_index else " "
                print(f" {prefix} {i}. {opt}")
            raw = input(f"请输入序号 (默认 {default_index+1}): ").strip()
            if raw == "":
                return options[default_index]
            if raw.lower() == 'q':
                return None
            if raw.isdigit():
                idx = int(raw) - 1
                if 0 <= idx < len(options):
                    return options[idx]
            # 文本匹配
            for opt in options:
                if raw.lower() == opt.lower():
                    return opt
            print("无效输入，请重试。")
    except KeyboardInterrupt:
        print("\n操作已取消")
        return options[default_index]

def make_safe_title(title: str) -> str:
    """根据合成命名规则，生成安全的标题前缀。"""
    safe_title = (
        title.replace(' ', '_')
             .replace('/', '_')
             .replace('\\', '_')
             .replace(':', '_')
             .replace('?', '_')
             .replace('*', '_')
             .replace('"', '_')
             .replace('<', '_')
             .replace('>', '_')
             .replace('|', '_')
    )
    return safe_title

def validate_media_assets(script_data: Dict[str, Any], images_dir: str, voice_dir: str) -> Dict[str, Any]:
    """校验图片、音频与脚本段落是否匹配，及命名规范。
    
    要求：
    - 图片: segment_1.png...segment_N.png 连续且齐全
    - 音频: voice_1.(wav|mp3)...voice_N.(wav|mp3) 连续且齐全
    - 数量与 script_data['segments'] 一致
    """
    issues: List[str] = []
    segments = script_data.get('segments', [])
    num_segments = len(segments)
    # 不再依赖标题作为音频命名前缀

    # 收集实际文件
    try:
        image_files = [f for f in os.listdir(images_dir) if os.path.isfile(os.path.join(images_dir, f))]
    except Exception:
        image_files = []
    try:
        audio_files = [f for f in os.listdir(voice_dir) if os.path.isfile(os.path.join(voice_dir, f))]
    except Exception:
        audio_files = []

    # 解析编号
    image_indices: List[int] = []
    for f in image_files:
        m = re.match(r'^segment_(\d+)\.(png|jpg|jpeg)$', f, re.IGNORECASE)
        if m:
            image_indices.append(int(m.group(1)))
    audio_indices: List[int] = []
    for f in audio_files:
        m = re.match(r'^voice_(\d+)\.(wav|mp3)$', f)
        if m:
            audio_indices.append(int(m.group(1)))

    # 基础数量检查
    if len(image_indices) != num_segments:
        issues.append(f"图片数量不匹配：期望{num_segments}张，实际{len(image_indices)}张")
    if len(audio_indices) != num_segments:
        issues.append(f"音频数量不匹配：期望{num_segments}段，实际{len(audio_indices)}段")

    # 连续性检查（1..N）
    expected_set = set(range(1, num_segments + 1))
    missing_images = sorted(list(expected_set - set(image_indices)))
    extra_images = sorted(list(set(image_indices) - expected_set))
    if missing_images:
        issues.append(f"缺少图片: segment_{missing_images[0]}...（共{len(missing_images)}个缺口）")
    if extra_images:
        issues.append(f"存在多余图片编号: {extra_images}")

    missing_audio = sorted(list(expected_set - set(audio_indices)))
    extra_audio = sorted(list(set(audio_indices) - expected_set))
    if missing_audio:
        issues.append(f"缺少音频: voice_{missing_audio[0]}.*（共{len(missing_audio)}个缺口）")
    if extra_audio:
        issues.append(f"存在多余音频编号: {extra_audio}")

    ok = len(issues) == 0
    return {
        'ok': ok,
        'issues': issues,
        'images_dir': images_dir,
        'voice_dir': voice_dir,
        'num_segments': num_segments
    }

class ProgressTracker:
    """进度跟踪器"""
    
    def __init__(self, total_steps: int, operation_name: str = "处理"):
        self.total_steps = total_steps
        self.current_step = 0
        self.operation_name = operation_name
        self.start_time = datetime.datetime.now()
    
    def step(self, message: str = ""):
        """前进一步"""
        self.current_step += 1
        progress = (self.current_step / self.total_steps) * 100
        
        elapsed = (datetime.datetime.now() - self.start_time).total_seconds()
        if self.current_step > 0 and elapsed > 0:
            eta = elapsed * (self.total_steps - self.current_step) / self.current_step
            eta_str = f"预计剩余: {eta:.0f}秒"
        else:
            eta_str = ""
        
        log_msg = f"{self.operation_name}进度: {self.current_step}/{self.total_steps} ({progress:.1f}%)"
        if message:
            log_msg += f" - {message}"
        if eta_str:
            log_msg += f" - {eta_str}"
        
        logger.info(log_msg)
    
    def complete(self):
        """标记完成"""
        total_time = (datetime.datetime.now() - self.start_time).total_seconds()
        logger.info(f"{self.operation_name}完成！总耗时: {total_time:.1f}秒")

def scan_input_files(input_dir: str = "input") -> List[Dict[str, Any]]:
    """
    扫描input文件夹中的PDF、EPUB和MOBI文件
    
    Args:
        input_dir: 输入文件夹路径
    
    Returns:
        List[Dict[str, Any]]: 文件信息列表，包含路径、名称、大小等信息
    """
    # 将相对路径锚定到项目目录（本文件所在目录）
    if not os.path.isabs(input_dir):
        input_dir = os.path.join(os.path.dirname(__file__), input_dir)
    
    if not os.path.exists(input_dir):
        logger.warning(f"输入目录不存在: {input_dir}")
        return []
    
    supported_extensions = ['.pdf', '.epub', '.mobi']
    files = []
    
    logger.info(f"正在扫描 {input_dir} 文件夹...")
    
    try:
        for file_name in os.listdir(input_dir):
            file_path = os.path.join(input_dir, file_name)
            
            # 跳过目录
            if os.path.isdir(file_path):
                continue
            
            # 检查文件扩展名
            file_extension = Path(file_path).suffix.lower()
            if file_extension in supported_extensions:
                file_info = get_file_info(file_path)
                files.append(file_info)
                logger.debug(f"找到文件: {file_name} ({file_info['size_formatted']})")
    
    except Exception as e:
        logger.error(f"扫描文件夹失败: {str(e)}")
        raise FileProcessingError(f"扫描文件夹失败: {str(e)}")
    
    # 按修改时间排序，最新的在前
    files.sort(key=lambda x: x['modified_time'], reverse=True)
    
    pdf_count = sum(1 for f in files if f['extension'] == '.pdf')
    epub_count = sum(1 for f in files if f['extension'] == '.epub')
    mobi_count = sum(1 for f in files if f['extension'] == '.mobi')
    logger.info(f"共找到 {len(files)} 个文件 (PDF: {pdf_count}, EPUB: {epub_count}, MOBI: {mobi_count})")
    
    return files

def display_file_menu(files: List[Dict[str, Any]]) -> None:
    """
    显示文件选择菜单
    
    Args:
        files: 文件信息列表
    """
    print("\n" + "="*60)
    print("📚 发现以下可处理的文件:")
    print("="*60)
    
    if not files:
        print("❌ 在input文件夹中未找到PDF、EPUB或MOBI文件")
        print("请将要处理的PDF、EPUB或MOBI文件放入input文件夹中")
        return
    
    for i, file_info in enumerate(files, 1):
        if file_info['extension'] == '.epub':
            file_type = "📖 EPUB"
        elif file_info['extension'] == '.pdf':
            file_type = "📄 PDF"
        elif file_info['extension'] == '.mobi':
            file_type = "📱 MOBI"
        else:
            file_type = "📄 FILE"
        modified_date = file_info['modified_time'].strftime('%Y-%m-%d %H:%M')
        
        print(f"{i:2}. {file_type} {file_info['name']}")
        print(f"     大小: {file_info['size_formatted']} | 修改时间: {modified_date}")
        print()

def get_user_file_selection(files: List[Dict[str, Any]]) -> Optional[str]:
    """
    获取用户的文件选择
    
    Args:
        files: 文件信息列表
    
    Returns:
        Optional[str]: 选择的文件路径，如果用户取消则返回None
    """
    if not files:
        return None
    
    while True:
        try:
            print("="*60)
            choice = input(f"请选择要处理的文件 (1-{len(files)}) 或输入 'q' 返回上一级: ").strip()
            
            if choice.lower() == 'q':
                print("👋 返回上一级")
                return None
            
            file_index = int(choice) - 1
            
            if 0 <= file_index < len(files):
                selected_file = files[file_index]
                print(f"\n✅ 您选择了: {selected_file['name']}")
                print(f"   文件大小: {selected_file['size_formatted']}")
                print(f"   文件类型: {selected_file['extension'].upper()}")
                # 直接返回所选文件路径，无需再次确认
                return selected_file['path']
            else:
                print(f"❌ 无效选择，请输入 1-{len(files)} 之间的数字")
                
        except ValueError:
            print("❌ 请输入有效的数字")
        except KeyboardInterrupt:
            print("\n\n👋 程序已取消")
            return None

def interactive_file_selector(input_dir: str = "input") -> Optional[str]:
    """
    交互式文件选择器
    
    Args:
        input_dir: 输入文件夹路径
    
    Returns:
        Optional[str]: 选择的文件路径，如果用户取消则返回None
    """
    print("\n🚀 智能视频制作系统")
    print("正在扫描可处理的文件...")
    
    # 扫描文件
    files = scan_input_files(input_dir)
    
    # 显示菜单
    display_file_menu(files)
    
    # 获取用户选择
    return get_user_file_selection(files)

# =============================
# 项目管理与进度检测（output/）
# =============================

def scan_output_projects(output_dir: str = "output") -> List[Dict[str, Any]]:
    """
    扫描 output 目录下的项目文件夹（约定：文件夹内包含 images/ voice/ text/ 等子目录）。

    Returns:
        List[Dict]: 每个项目的 { path, name, modified_time } 信息
    """
    # 将相对路径锚定到项目目录（本文件所在目录）
    if not os.path.isabs(output_dir):
        output_dir = os.path.join(os.path.dirname(__file__), output_dir)

    projects: List[Dict[str, Any]] = []
    if not os.path.exists(output_dir):
        return projects

    try:
        for entry in os.listdir(output_dir):
            p = os.path.join(output_dir, entry)
            if not os.path.isdir(p):
                continue
            # 粗略判断：包含 text/ 目录即认为是项目
            text_dir = os.path.join(p, "text")
            if os.path.isdir(text_dir):
                stat = os.stat(p)
                projects.append({
                    "path": p,
                    "name": entry,
                    "modified_time": datetime.datetime.fromtimestamp(stat.st_mtime)
                })
    except Exception as e:
        logger.warning(f"扫描输出目录失败: {e}")
        return []

    # 最新修改在前
    projects.sort(key=lambda x: x["modified_time"], reverse=True)
    return projects

def display_project_menu(projects: List[Dict[str, Any]]) -> None:
    print("\n" + "="*60)
    print("📂 发现以下现有项目:")
    print("="*60)
    if not projects:
        print("❌ 在 output 目录中未找到现有项目")
        return
    for i, info in enumerate(projects, 1):
        modified_date = info['modified_time'].strftime('%Y-%m-%d %H:%M')
        print(f"{i:2}. {info['name']}")
        print(f"     修改时间: {modified_date}")
        print()

def get_user_project_selection(projects: List[Dict[str, Any]]) -> Optional[str]:
    if not projects:
        return None
    while True:
        try:
            print("="*60)
            choice = input(f"请选择要打开的项目 (1-{len(projects)}) 或输入 'q' 返回上一级: ").strip()
            if choice.lower() == 'q':
                return None
            idx = int(choice) - 1
            if 0 <= idx < len(projects):
                selected = projects[idx]
                print(f"\n✅ 您选择了项目: {selected['name']}")
                return selected['path']
            else:
                print(f"❌ 无效选择，请输入 1-{len(projects)} 之间的数字")
        except ValueError:
            print("❌ 请输入有效的数字")
        except KeyboardInterrupt:
            print("\n操作已取消")
            return None

def interactive_project_selector(output_dir: str = "output") -> Optional[str]:
    """
    交互式项目选择器（从 output/ 选择已有项目文件夹）
    """
    print("\n📂 打开现有项目")
    print("正在扫描 output 目录...")
    projects = scan_output_projects(output_dir)
    display_project_menu(projects)
    return get_user_project_selection(projects)

def _read_json_if_exists(path: str) -> Optional[Dict[str, Any]]:
    try:
        if os.path.exists(path):
            with open(path, 'r', encoding='utf-8') as f:
                return json.load(f)
    except Exception as e:
        logger.warning(f"读取JSON失败 {path}: {e}")
    return None

def detect_project_progress(project_dir: str) -> Dict[str, Any]:
    """
    检测项目当前进度阶段。

    Returns:
        一个进度字典，其中：
        - current_step: 内部步骤编号（2..6），对应 5 步展示的 1..5（= current_step-1）
        - current_step_display: 对用户展示的步骤编号（1..5）
        其余布尔标记用于判定各阶段产物是否就绪。
    """
    text_dir = os.path.join(project_dir, "text")
    images_dir = os.path.join(project_dir, "images")
    voice_dir = os.path.join(project_dir, "voice")
    final_video_path = os.path.join(project_dir, "final_video.mp4")

    script = _read_json_if_exists(os.path.join(text_dir, "script.json"))
    has_script = script is not None and isinstance(script, dict) and 'segments' in script

    keywords = _read_json_if_exists(os.path.join(text_dir, "keywords.json"))
    has_keywords = has_script and keywords is not None and 'segments' in keywords and \
        len(keywords.get('segments', [])) == len(script.get('segments', []))

    images_ok = False
    audio_ok = False
    if has_script:
        try:
            v = validate_media_assets(script_data=script, images_dir=images_dir, voice_dir=voice_dir)
            # 只看图片或音频是否分别就绪
            # 图片就绪: 没有图片数量/连续性问题
            # 音频就绪: 没有音频数量/连续性问题
            # 简化：先通过两次局部检查
            # 图片检查
            num_segments = len(script.get('segments', []))
            image_files = [f for f in os.listdir(images_dir) if os.path.isfile(os.path.join(images_dir, f))] if os.path.isdir(images_dir) else []
            import re as _re
            image_indices = []
            for f in image_files:
                m = _re.match(r'^segment_(\d+)\.(png|jpg|jpeg)$', f, _re.IGNORECASE)
                if m:
                    image_indices.append(int(m.group(1)))
            images_ok = (len(image_indices) == num_segments) and (set(image_indices) == set(range(1, num_segments+1)))
            # 音频检查
            audio_files = [f for f in os.listdir(voice_dir) if os.path.isfile(os.path.join(voice_dir, f))] if os.path.isdir(voice_dir) else []
            audio_indices = []
            for f in audio_files:
                m = _re.match(r'^voice_(\d+)\.(wav|mp3)$', f)
                if m:
                    audio_indices.append(int(m.group(1)))
            audio_ok = (len(audio_indices) == num_segments) and (set(audio_indices) == set(range(1, num_segments+1)))
        except Exception:
            images_ok = False
            audio_ok = False

    has_final_video = os.path.exists(final_video_path) and os.path.getsize(final_video_path) > 0

    # 计算 current_step（对内与对外一致：1..5）
    current_step = 0
    if has_script:
        current_step = 1
    if has_keywords:
        current_step = 2
    if images_ok:
        current_step = 3
    if audio_ok:
        current_step = 4
    if has_final_video:
        current_step = 5

    return {
        'has_script': has_script,
        'has_keywords': has_keywords,
        'images_ok': images_ok,
        'audio_ok': audio_ok,
        'has_final_video': has_final_video,
        'current_step': current_step,
        'current_step_display': max(1, min(5, current_step)),
        'script': script,
        'keywords': keywords,
        'final_video_path': final_video_path,
        'images_dir': images_dir,
        'voice_dir': voice_dir,
        'text_dir': text_dir
    }

from typing import Optional

def prompt_step_to_rerun(current_step: int) -> Optional[int]:
    """
    询问用户要从哪一步开始重做（展示 1..5）。
    - 输入 1..5（展示层）将映射到内部 2..6（核心逻辑层）。
    - 返回值为内部步骤编号（2..6）；输入 q/CTRL-C 返回 None。
    """
    # 对外展示 1..5（合并了文档读取+智能缩写），但内部仍映射到 2..6
    options = [
        "第1步：智能缩写",
        "第2步：关键词提取",
        "第3步：AI图像生成",
        "第4步：语音合成",
        "第5步：视频合成",
    ]
    # 统一 1..5（对内对外一致）
    current_display_step = max(1, min(5, current_step))
    print("\n当前项目进度（共5步）：已完成到第{}步".format(current_display_step))
    for i, opt in enumerate(options, 1):
        marker = '*' if i == current_display_step else ' '
        print(f" {marker} {i}. {opt}")
    default_display = current_display_step
    while True:
        try:
            raw = input(f"请输入步骤号 1-5 或输入 'q' 返回上一级 (默认 {default_display}): ").strip()
            if raw == "":
                return default_display
            if raw.lower() == 'q':
                return None
            if raw.isdigit():
                n = int(raw)
                if 1 <= n <= 5:
                    return n
            print("无效输入，请输入 1-5。")
        except KeyboardInterrupt:
            print("\n操作已取消")
            return None

def collect_ordered_assets(project_dir: str, script_data: Dict[str, Any], require_audio: bool = True) -> Dict[str, List[str]]:
    """
    根据 script_data 的段落顺序，收集按序排列的图片和（可选）音频文件路径。

    Args:
        project_dir: 项目目录
        script_data: 包含段落信息的脚本数据
        require_audio: 是否强制要求每段音频都存在；为 False 时仅收集图片，音频如存在则收集，不存在不报错。

    Returns:
        Dict[str, List[str]]: {"images": [...], "audio": [...]}；当 require_audio=False 且音频不存在时，"audio" 可为空列表。
    """
    images_dir = os.path.join(project_dir, "images")
    voice_dir = os.path.join(project_dir, "voice")
    # 不再依赖标题作为音频命名前缀
    num_segments = len(script_data.get('segments', []))

    image_paths: List[str] = []
    audio_paths: List[str] = []
    for i in range(1, num_segments+1):
        # 按多种常见图片后缀依次探测
        _candidates = [
            os.path.join(images_dir, f"segment_{i}.png"),
            os.path.join(images_dir, f"segment_{i}.jpg"),
            os.path.join(images_dir, f"segment_{i}.jpeg"),
        ]
        image_path = None
        for _p in _candidates:
            if os.path.exists(_p):
                image_path = _p
                break
        audio_wav_new = os.path.join(voice_dir, f"voice_{i}.wav")
        audio_mp3_new = os.path.join(voice_dir, f"voice_{i}.mp3")
        if not image_path:
            # 以 .png 为主的规范名称提示
            raise FileNotFoundError(f"缺少图片: segment_{i}.(png|jpg|jpeg)")
        if require_audio:
            if os.path.exists(audio_wav_new):
                audio_path = audio_wav_new
            elif os.path.exists(audio_mp3_new):
                audio_path = audio_mp3_new
            else:
                raise FileNotFoundError(f"缺少音频: voice_{i}.(wav|mp3)")
            audio_paths.append(audio_path)
        else:
            # 非强制音频：有则收集，无则跳过
            if os.path.exists(audio_wav_new):
                audio_paths.append(audio_wav_new)
            elif os.path.exists(audio_mp3_new):
                audio_paths.append(audio_mp3_new)
        image_paths.append(image_path)
    return {"images": image_paths, "audio": audio_paths}

def clear_downstream_outputs(project_dir: str, from_step: int) -> None:
    """
    清理从指定步骤之后的产物，以便重新生成。
    from_step: 1..5（对外/对内统一步骤编号）
    """
    text_dir = os.path.join(project_dir, "text")
    images_dir = os.path.join(project_dir, "images")
    voice_dir = os.path.join(project_dir, "voice")
    final_video_path = os.path.join(project_dir, "final_video.mp4")

    try:
        if from_step <= 1:
            # 删除 keywords
            kp = os.path.join(text_dir, "keywords.json")
            if os.path.exists(kp):
                os.remove(kp)
        if from_step <= 2:
            # 清空 images
            if os.path.isdir(images_dir):
                for f in os.listdir(images_dir):
                    fp = os.path.join(images_dir, f)
                    if os.path.isfile(fp):
                        os.remove(fp)
        if from_step <= 3:
            # 清空 voice
            if os.path.isdir(voice_dir):
                for f in os.listdir(voice_dir):
                    fp = os.path.join(voice_dir, f)
                    if os.path.isfile(fp):
                        os.remove(fp)
        if from_step <= 4:
            # 删除最终视频
            if os.path.exists(final_video_path):
                os.remove(final_video_path)
    except Exception as e:
        logger.warning(f"清理旧产物失败: {e}")

def export_script_to_docx(script_data: Dict[str, Any], docx_path: str) -> str:
    """
    将脚本JSON导出为可阅读的DOCX文档，仅包含标题与各段content。

    要求：
    - 1.5倍行距
    - 字体：宋体（含东亚字体设置）
    - 正文两端对齐，标题居中

    Args:
        script_data: 含有 title 与 segments 的脚本数据
        docx_path: 输出的docx文件完整路径（建议位于 output/{project}/text/script.docx ）

    Returns:
        str: 实际保存的docx路径
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_LINE_SPACING
    from docx.oxml.ns import qn

    # 创建文档
    document = Document()

    # 尝试全局设置 Normal 样式字体与行距
    try:
        normal_style = document.styles['Normal']
        normal_style.font.name = '宋体'
        # 设置东亚字体
        if hasattr(normal_style, 'element') and normal_style.element is not None:
            rPr = normal_style.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        # 行距 1.5 倍
        if normal_style.paragraph_format is not None:
            try:
                normal_style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            except Exception:
                normal_style.paragraph_format.line_spacing = 1.5
    except Exception:
        # 即使全局样式设置失败，也不影响后续逐段设置
        pass

    title_text = script_data.get('title', 'untitled')
    segments = script_data.get('segments', []) or []

    # 标题（居中）
    title_para = document.add_paragraph()
    title_run = title_para.add_run(title_text)
    title_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    # 标题字体（宋体）
    try:
        title_run.font.name = '宋体'
        title_run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass
    # 标题行距 1.5
    try:
        title_para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except Exception:
        title_para.paragraph_format.line_spacing = 1.5

    # 正文段落（两端对齐，每个content独立段落）
    for seg in segments:
        content = (seg or {}).get('content', '')
        if not content:
            continue
        p = document.add_paragraph()
        r = p.add_run(content)
        # 字体宋体（含东亚）
        try:
            r.font.name = '宋体'
            r._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        except Exception:
            pass
        # 段落两端对齐 + 1.5倍行距
        p.alignment = WD_ALIGN_PARAGRAPH.JUSTIFY
        try:
            p.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
        except Exception:
            p.paragraph_format.line_spacing = 1.5

    ensure_directory_exists(os.path.dirname(docx_path))
    document.save(docx_path)
    logger.info(f"阅读版DOCX已保存: {docx_path}")
    return docx_path

# 导出主要函数和类
__all__ = [
    'VideoProcessingError', 'APIError', 'FileProcessingError',
    'log_function_call', 'ensure_directory_exists', 'clean_text',
    'validate_file_format', 'safe_json_loads', 'save_json_file', 'load_json_file',
    'calculate_duration', 'format_file_size', 'get_file_info',
    'retry_on_failure', 'validate_required_fields', 'create_processing_summary',
    'progress_callback', 'ProgressTracker', 'logger',
    'scan_input_files', 'display_file_menu', 'get_user_file_selection', 'interactive_file_selector',
    'scan_output_projects', 'interactive_project_selector', 'detect_project_progress', 'prompt_step_to_rerun',
    'collect_ordered_assets', 'clear_downstream_outputs', 'display_project_menu', 'get_user_project_selection',
    'export_script_to_docx',
]