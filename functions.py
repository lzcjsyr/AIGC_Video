"""
智能视频制作系统 - 核心功能模块
包含文档读取、智能处理、图像生成、语音合成、视频制作等功能
"""

from moviepy import ImageClip, AudioFileClip, CompositeVideoClip, concatenate_videoclips, TextClip, ColorClip, CompositeAudioClip
# MoviePy 2.x: 使用类效果 API
try:
    from moviepy.audio.fx.AudioLoop import AudioLoop  # type: ignore
except Exception:
    AudioLoop = None  # fallback later
from moviepy.audio.fx.MultiplyVolume import MultiplyVolume  # type: ignore
from typing import Optional, Dict, Any, List, Tuple
import requests
import json
import os
import re
import datetime
import ebooklib
from ebooklib import epub
import PyPDF2
import pdfplumber
try:
    import fitz  # pymupdf
    FITZ_AVAILABLE = True
except ImportError:
    FITZ_AVAILABLE = False

from prompts import summarize_system_prompt, keywords_extraction_prompt, IMAGE_STYLE_PRESETS, OPENING_IMAGE_STYLES
from genai_api import text_to_text, text_to_image_doubao, text_to_audio_bytedance
from config import config
 
from utils import (
    logger, FileProcessingError, APIError,
    log_function_call, ensure_directory_exists, clean_text, 
    validate_file_format
)
from utils import parse_json_robust
import numpy as np

# 统一字体解析：优先使用系统中文字体路径，失败回退到传入名称
def resolve_font_path(preferred: Optional[str]) -> Optional[str]:
    if preferred and os.path.exists(preferred):
        return preferred
    candidate_paths = [
        "/System/Library/Fonts/PingFang.ttc",
        "/System/Library/Fonts/Hiragino Sans GB.ttc",
        "/System/Library/Fonts/Supplemental/Songti.ttc",
        "/System/Library/Fonts/STHeiti Light.ttc",
        "/System/Library/Fonts/Supplemental/SimHei.ttf",
        "/System/Library/Fonts/Supplemental/SimSun.ttf",
        "/Library/Fonts/Arial Unicode.ttf",
        "/Library/Fonts/Arial Unicode MS.ttf",
    ]
    for path in candidate_paths:
        if os.path.exists(path):
            return path
    return preferred

# 通用下载器：下载二进制内容并保存到指定路径
def download_to_path(url: str, output_path: str, error_msg: str = "下载失败") -> None:
    response = requests.get(url)
    if response.status_code != 200:
        raise ValueError(error_msg)
    with open(output_path, 'wb') as f:
        f.write(response.content)

################ Document Reading ################
@log_function_call
def read_document(file_path: str) -> Tuple[str, int]:
    """
    读取EPUB、PDF、MOBI、DOCX、DOC文档，返回内容和字数
    简化分发逻辑，统一清理与统计。
    """
    validate_file_format(file_path, config.SUPPORTED_INPUT_FORMATS)
    ext = os.path.splitext(file_path)[1].lower()
    logger.info(f"开始读取{ext.upper()}文件: {os.path.basename(file_path)}")

    if ext == '.epub':
        return read_epub(file_path)
    if ext == '.pdf':
        return read_pdf(file_path)
    if ext == '.mobi':
        return read_mobi(file_path)
    if ext == '.azw3':
        return read_azw3(file_path)
    if ext in ('.docx', '.doc'):
        return read_word(file_path)
    raise FileProcessingError(f"不支持的文件格式: {ext}")

def read_epub(file_path: str) -> Tuple[str, int]:
    """读取EPUB文件内容"""
    try:
        book = epub.read_epub(file_path)
        content_parts = []
        
        logger.debug("正在提取EPUB文件中的文本内容...")
        
        # 获取所有文本内容
        for item in book.get_items():
            if item.get_type() == ebooklib.ITEM_DOCUMENT:
                content = item.get_content().decode('utf-8')
                # 清理HTML标签和格式化文本
                content = clean_text(content)
                if content:
                    content_parts.append(content)
        
        if not content_parts:
            raise FileProcessingError("EPUB文件中未找到可读取的文本内容")
        
        full_content = ' '.join(content_parts)
        word_count = len(full_content)
        
        logger.info(f"EPUB文件读取成功，总字数: {word_count:,}字")
        return full_content, word_count
    
    except Exception as e:
        logger.error(f"读取EPUB文件失败: {str(e)}")
        raise FileProcessingError(f"读取EPUB文件失败: {str(e)}")

def read_pdf(file_path: str) -> Tuple[str, int]:
    """读取PDF文件内容"""
    try:
        content_parts = []
        
        # 首选：PyMuPDF（fitz）——速度快、鲁棒性好
        if FITZ_AVAILABLE:
            logger.debug("使用PyMuPDF提取PDF文本...")
            try:
                doc = fitz.open(file_path)
                for i in range(len(doc)):
                    page = doc.load_page(i)
                    text = page.get_text()
                    if text and text.strip():
                        content_parts.append(text)
                        logger.debug(f"已提取第{i+1}页内容，字符数: {len(text)}")
                doc.close()
            except Exception as e:
                logger.debug(f"PyMuPDF提取失败: {str(e)}")
        
        # 备用：pdfplumber（如未安装PyMuPDF或提取为空）
        if not content_parts:
            logger.debug("尝试使用pdfplumber提取PDF文本...")
            with pdfplumber.open(file_path) as pdf:
                for i, page in enumerate(pdf.pages, 1):
                    text = page.extract_text()
                    if text:
                        content_parts.append(text)
                        logger.debug(f"已提取第{i}页内容，字符数: {len(text)}")
        
        if not content_parts:
            raise FileProcessingError("无法从PDF文件中提取文本内容，可能是扫描版PDF或编码问题")
        
        full_content = ' '.join(content_parts)
        
        # 清理文本
        cleaned_content = clean_text(full_content)
        cleaned_length = len(cleaned_content)
        
        # 分析内容质量
        cid_count = full_content.count('(cid:')
        total_chars = len(full_content)
        chinese_chars = sum(1 for c in cleaned_content if '\u4e00' <= c <= '\u9fff')
        english_chars = sum(1 for c in cleaned_content if c.isalpha() and ord(c) < 128)
        readable_chars = chinese_chars + english_chars
        
        # 计算损坏程度
        loss_ratio = (total_chars - cleaned_length) / max(1, total_chars)  # 清理时丢失的字符比例
        readable_ratio = readable_chars / max(1, cleaned_length)  # 清理后的可读比例
        
        logger.info(f"PDF内容分析: 原始={total_chars:,}, 清理后={cleaned_length:,}, 可读={readable_chars:,}")
        logger.info(f"质量评估: 内容丢失={loss_ratio:.1%}, 可读性={readable_ratio:.1%}")
        
        # 严格的乱码检测：避免将大量乱码发送给LLM
        if (loss_ratio > 0.8 or  # 超过80%的内容在清理时丢失
            readable_ratio < 0.5 or  # 可读字符低于50%
            cleaned_length < 1000 or  # 清理后内容过少
            readable_chars < 5000):    # 可读字符绝对数量过少
            
            # 判断主要问题类型
            if cid_count > total_chars * 0.1:
                main_issue = "CID字体编码问题"
                issue_detail = f"检测到 {cid_count:,} 个CID字符 ({cid_count/total_chars:.1%})"
            elif loss_ratio > 0.8:
                main_issue = "字体映射损坏"
                issue_detail = f"清理时丢失了 {loss_ratio:.1%} 的内容，大部分为乱码字符"
            else:
                main_issue = "内容质量不足"
                issue_detail = f"可读内容太少，只有 {readable_chars:,} 个有效字符"
            
            logger.error(f"PDF内容质量不佳：{main_issue}（{issue_detail}）")
            raise FileProcessingError(
                "PDF文本质量不足，可能为扫描版或编码异常。建议使用OCR识别，或先转换为EPUB/TXT后重试。"
            )
        
        logger.info(f"PDF内容检测通过: 中文字符={chinese_chars}, 英文字符={english_chars}, 总长度={cleaned_length}")
        logger.info(f"PDF文件读取成功，总字数: {cleaned_length:,}字")
        return cleaned_content, cleaned_length
    
    except FileProcessingError:
        # 直接重新抛出FileProcessingError，避免重复包装
        raise
    except Exception as e:
        logger.error(f"读取PDF文件失败: {str(e)}")
        raise FileProcessingError(f"读取PDF文件失败: {str(e)}")

def read_mobi(file_path: str) -> Tuple[str, int]:
    """读取MOBI文件内容"""
    import struct
    import re
    
    try:
        logger.debug("正在读取MOBI文件...")
        
        with open(file_path, 'rb') as f:
            # 读取文件头，检查是否为MOBI格式
            f.seek(60)  # MOBI标识符位置
            mobi_header = f.read(4)
            
            if mobi_header != b'MOBI':
                # 如果不是标准MOBI，尝试作为AZW格式
                f.seek(0)
                content = f.read()
                if b'BOOKMOBI' not in content[:100]:
                    raise FileProcessingError("文件不是有效的MOBI格式")
            
            # 简单的MOBI文本提取
            f.seek(0)
            raw_content = f.read()
            
            # 查找文本内容区域
            text_content = ""
            
            # 方法1: 搜索HTML标签内容
            html_matches = re.findall(b'<.*?>(.*?)</.*?>', raw_content, re.DOTALL)
            for match in html_matches:
                try:
                    text = match.decode('utf-8', errors='ignore')
                    if text.strip():
                        text_content += text + " "
                except:
                    continue
            
            # 方法2: 如果没有找到HTML，尝试直接提取可读文本
            if not text_content.strip():
                try:
                    # 寻找文本记录的开始位置
                    for i in range(0, len(raw_content) - 100, 1000):
                        chunk = raw_content[i:i+1000]
                        try:
                            decoded = chunk.decode('utf-8', errors='ignore')
                            # 过滤出可能的文本内容
                            clean_chunk = re.sub(r'[^\w\s\u4e00-\u9fff\u3000-\u303f\uff00-\uffef.,!?;:"\'()[\]{}]', ' ', decoded)
                            if len(clean_chunk.strip()) > 50:  # 如果有足够的可读内容
                                words = clean_chunk.split()
                                # 过滤掉过短的"词"
                                valid_words = [w for w in words if len(w) >= 2 or re.match(r'[\u4e00-\u9fff]', w)]
                                if len(valid_words) > 5:
                                    text_content += " ".join(valid_words) + " "
                        except:
                            continue
                except:
                    pass
            
            if not text_content.strip():
                raise FileProcessingError("无法从MOBI文件中提取文本内容，可能是加密文件或格式不支持")
            
            # 清理提取的文本
            cleaned_content = clean_text(text_content)
            word_count = len(cleaned_content)
            
            # 质量检查
            if word_count < 100:
                raise FileProcessingError("MOBI文件内容过少，可能是加密文件或解析失败")
            
            logger.info(f"MOBI文件读取成功，总字数: {word_count:,}字")
            return cleaned_content, word_count
    
    except Exception as e:
        logger.error(f"读取MOBI文件失败: {str(e)}")
        raise FileProcessingError(f"读取MOBI文件失败: {str(e)}。建议使用Calibre转换为EPUB格式后重试。")

def read_azw3(file_path: str) -> Tuple[str, int]:
    """
    读取AZW3文件内容
    AZW3是Amazon Kindle的专有格式，基于MOBI演进，支持HTML5/CSS3
    """
    try:
        logger.debug("正在读取AZW3文件...")
        
        # 首先尝试使用mobi库解析（最佳方案）
        try:
            import mobi
            import tempfile
            import shutil
            
            # 使用mobi库提取内容
            tempdir, filepath = mobi.extract(file_path)
            
            if not filepath or not os.path.exists(filepath):
                raise FileProcessingError("mobi库解析失败：未生成有效的输出文件")
            
            # 根据输出文件类型读取内容
            ext = os.path.splitext(filepath)[1].lower()
            text_content = ""
            
            if ext == '.epub':
                # 如果输出为EPUB，使用现有的EPUB读取逻辑
                content, word_count = read_epub(filepath)
                # 清理临时文件
                try:
                    shutil.rmtree(tempdir)
                except:
                    pass
                logger.info(f"AZW3文件读取成功（通过mobi库转EPUB），总字数: {word_count:,}字")
                return content, word_count
            
            elif ext in ['.html', '.htm']:
                # 如果输出为HTML，解析HTML内容
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    html_content = f.read()
                
                # 简单的HTML标签清理
                import re
                # 移除script和style标签及其内容
                html_content = re.sub(r'<(script|style)[^>]*>.*?</\1>', '', html_content, flags=re.DOTALL | re.IGNORECASE)
                # 移除HTML标签，保留文本
                html_content = re.sub(r'<[^>]+>', '', html_content)
                # 解码HTML实体
                import html
                text_content = html.unescape(html_content)
                
            elif ext == '.txt':
                # 如果输出为纯文本
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
            
            else:
                # 尝试直接读取为文本
                with open(filepath, 'r', encoding='utf-8', errors='ignore') as f:
                    text_content = f.read()
            
            # 清理临时文件
            try:
                shutil.rmtree(tempdir)
            except:
                pass
            
            if not text_content.strip():
                raise FileProcessingError("mobi库解析成功但未提取到文本内容")
            
            # 清理提取的文本
            cleaned_content = clean_text(text_content)
            word_count = len(cleaned_content)
            
            # 质量检查
            if word_count < 100:
                raise FileProcessingError("AZW3文件内容过少，可能是加密文件或解析失败")
            
            logger.info(f"AZW3文件读取成功（通过mobi库），总字数: {word_count:,}字")
            return cleaned_content, word_count
            
        except ImportError:
            logger.warning("mobi库未安装，尝试回退到基础MOBI解析")
            # 回退到现有的MOBI解析逻辑
            return read_mobi(file_path)
            
        except Exception as e:
            logger.warning(f"mobi库解析失败: {str(e)}，尝试回退到基础MOBI解析")
            # 回退到现有的MOBI解析逻辑
            try:
                return read_mobi(file_path)
            except Exception as fallback_e:
                raise FileProcessingError(
                    f"AZW3文件解析失败。主要尝试（mobi库）: {str(e)}；"
                    f"回退尝试（基础解析）: {str(fallback_e)}。"
                    f"建议: 1) 确认文件无DRM保护；2) 使用Calibre转换为EPUB格式后重试"
                )
    
    except FileProcessingError:
        # 直接重新抛出FileProcessingError，避免重复包装
        raise
    except Exception as e:
        logger.error(f"读取AZW3文件失败: {str(e)}")
        raise FileProcessingError(f"读取AZW3文件失败: {str(e)}。建议使用Calibre转换为EPUB格式后重试。")

def read_word(file_path: str) -> Tuple[str, int]:
    """
    简洁的通用方案：
    - .docx: 直接用 python-docx 解析
    - .doc: 尝试用 LibreOffice 将 .doc 转 .docx 后再解析；失败则用 antiword 提取纯文本
    """
    import subprocess
    import tempfile
    ext = os.path.splitext(file_path)[1].lower()

    # 首先处理 .docx：纯 Python，最通用
    if ext == '.docx':
        try:
            from docx import Document
            doc = Document(file_path)
            parts: List[str] = []
            for p in doc.paragraphs:
                t = (p.text or '').strip()
                if t:
                    parts.append(t)
            for table in getattr(doc, 'tables', []) or []:
                for row in table.rows:
                    for cell in row.cells:
                        t = (cell.text or '').strip()
                        if t:
                            parts.append(t)
            raw = "\n".join(parts)
            cleaned = clean_text(raw)
            return cleaned, len(cleaned)
        except Exception as e:
            logger.error(f"读取DOCX失败: {e}")
            raise FileProcessingError("读取DOCX失败：请确认文件未损坏，或转为PDF/EPUB后重试。")

    # 处理 .doc：先尝试用 LibreOffice 转换为 .docx，再用 python-docx 解析；否则退回 antiword 取纯文本
    try:
        with tempfile.TemporaryDirectory() as tmpd:
            out_docx = os.path.join(tmpd, "converted.docx")
            # 先转 docx
            try:
                subprocess.run([
                    "soffice", "--headless", "--convert-to", "docx", file_path, "--outdir", tmpd
                ], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
                # 找到输出的 docx（LibreOffice会保留原文件名）
                base = os.path.splitext(os.path.basename(file_path))[0]
                candidate = os.path.join(tmpd, base + ".docx")
                if os.path.exists(candidate):
                    out_docx = candidate
                # 用 python-docx 解析转换后的 docx
                from docx import Document
                doc = Document(out_docx)
                parts = []
                for p in doc.paragraphs:
                    t = (p.text or '').strip()
                    if t:
                        parts.append(t)
                for table in getattr(doc, 'tables', []) or []:
                    for row in table.rows:
                        for cell in row.cells:
                            t = (cell.text or '').strip()
                            if t:
                                parts.append(t)
                raw = "\n".join(parts)
                cleaned = clean_text(raw)
                return cleaned, len(cleaned)
            except Exception:
                pass
    except Exception:
        pass

    # 退回 antiword（输出纯文本）
    try:
        result = subprocess.run(["antiword", file_path], stdout=subprocess.PIPE, stderr=subprocess.PIPE, check=True)
        raw = result.stdout.decode('utf-8', errors='ignore')
        cleaned = clean_text(raw)
        return cleaned, len(cleaned)
    except Exception:
        pass

    raise FileProcessingError("读取DOC失败：请安装 LibreOffice 或 antiword，或先将DOC转换为DOCX/PDF后重试。")

def read_docx(file_path: str) -> Tuple[str, int]:
    """向后兼容：转调统一的 read_word。"""
    return read_word(file_path)

def read_doc(file_path: str) -> Tuple[str, int]:
    """向后兼容：转调统一的 read_word。"""
    return read_word(file_path)

################ Intelligent Summarization ################
def intelligent_summarize(server: str, model: str, content: str, target_length: int, num_segments: int) -> Dict[str, Any]:
    """
    智能缩写 - 第一次LLM处理
    新逻辑：LLM生成完整口播终稿（content），返回原始数据，不进行分段。
    """
    try:
        user_message = f"""请将以下内容智能压缩为约{target_length}字的口播终稿，不要分段：

原文内容：
{content}

要求：
1. 保持核心信息与清晰逻辑，语言适合口播
2. 输出完整终稿到 content 字段，勿做任何分段
3. 总字数控制在{target_length}字左右
"""

        output = text_to_text(
            server=server,
            model=model,
            prompt=user_message,
            system_message=summarize_system_prompt,
            max_tokens=4096,
            temperature=config.LLM_TEMPERATURE_SCRIPT
        )

        if output is None:
            raise ValueError("未能从 API 获取响应。")

        parsed = parse_json_robust(output)

        # 新格式强约束：必须包含 title / content，golden_quote 可选
        if "title" not in parsed or "content" not in parsed:
            raise ValueError("生成的 JSON 缺少必需字段：title 或 content")

        title = parsed.get("title", "untitled")
        golden_quote = parsed.get("golden_quote", "")
        full_text = (parsed.get("content") or "").strip()
        if not full_text:
            raise ValueError("生成的 content 为空")

        # 返回原始数据，不进行分段
        raw_data: Dict[str, Any] = {
            "title": title,
            "golden_quote": golden_quote,
            "content": full_text,
            "total_length": len(full_text),
            "target_segments": num_segments,
            "created_time": datetime.datetime.now().isoformat(),
            "model_info": {
                "llm_server": server,
                "llm_model": model,
                "generation_type": "raw_generation"
            }
        }

        return raw_data

    except json.JSONDecodeError:
        raise ValueError("解析 JSON 输出失败")
    except Exception as e:
        raise ValueError(f"智能缩写处理错误: {e}")

def process_raw_to_script(raw_data: Dict[str, Any], num_segments: int) -> Dict[str, Any]:
    """
    将原始数据处理为分段脚本数据。
    这是步骤1.5的核心功能，从raw数据生成最终的script数据。
    
    Args:
        raw_data: 包含title、golden_quote、content的原始数据
        num_segments: 目标分段数量
        
    Returns:
        Dict[str, Any]: 分段后的脚本数据，格式与原来的script.json相同
    """
    try:
        title = raw_data.get("title", "untitled")
        golden_quote = raw_data.get("golden_quote", "")
        full_text = raw_data.get("content", "").strip()
        
        if not full_text:
            raise ValueError("原始数据的 content 字段为空")

        # 代码分段：按重标点优先，尽量均衡为 num_segments 段
        segments_text = _split_text_into_segments(full_text, num_segments)

        # 汇总统计
        total_length = len(full_text)
        enhanced_data: Dict[str, Any] = {
            "title": title,
            "golden_quote": golden_quote,
            "total_length": total_length,
            "target_segments": num_segments,
            "actual_segments": len(segments_text),
            "created_time": datetime.datetime.now().isoformat(),
            "model_info": raw_data.get("model_info", {}),
            "segments": []
        }
        
        # 更新模型信息中的处理类型
        enhanced_data["model_info"]["generation_type"] = "script_generation"

        # 估算每段时长
        wpm = int(getattr(config, "SPEECH_SPEED_WPM", 300))
        for i, seg_text in enumerate(segments_text, 1):
            length_i = len(seg_text)
            estimated_duration = length_i / max(1, wpm) * 60
            enhanced_data["segments"].append({
                "index": i,
                "content": seg_text,
                "length": length_i,
                "estimated_duration": round(estimated_duration, 1)
            })

        return enhanced_data

    except Exception as e:
        raise ValueError(f"处理原始数据为脚本错误: {e}")

def _split_text_into_segments(full_text: str, num_segments: int) -> List[str]:
    """
    使用重标点（。？！；.!?\n）先切成句子，再在句子边界上均衡聚合为 num_segments 段。
    若句子不足，则字符级均分补齐，保证输出恰好 num_segments 段。
    """
    text = (full_text or "").strip()
    if num_segments <= 1 or len(text) == 0:
        return [text] if text else [""]

    # 1) 句子级切分：将标点（包括换行符）附着到前句
    raw_parts = re.split(r'([。！？；.!?\n])', text)
    sentences: List[str] = []
    i = 0
    while i < len(raw_parts):
        token = raw_parts[i]
        if token and token.strip():
            cur = token.strip()
            if i + 1 < len(raw_parts) and raw_parts[i + 1] and raw_parts[i + 1].strip() in '。！？；.!?\n':
                cur += raw_parts[i + 1].strip()
                i += 2
            else:
                i += 1
            sentences.append(cur)
        else:
            i += 1

    if not sentences:
        sentences = [text]

    # 2) 若句子数量 >= 段数：在句子边界上均衡聚合
    total_len = sum(len(s) for s in sentences)
    if len(sentences) >= num_segments:
        ideal = total_len / float(num_segments)
        cum = 0
        boundaries: List[int] = []  # 选取 num_segments-1 个边界（句索引之后）
        for idx, s in enumerate(sentences):
            prev = cum
            cum += len(s)
            target_k = len(boundaries) + 1
            threshold = ideal * target_k
            # 当累计长度跨过阈值，或已经接近阈值，就在此边界切分
            if prev < threshold <= cum or (abs(cum - threshold) <= len(s) // 2):
                if len(boundaries) < num_segments - 1:
                    boundaries.append(idx)
            if len(boundaries) >= num_segments - 1:
                break

        # 根据边界组装段落
        segments: List[str] = []
        start = 0
        for b in boundaries:
            segment_text = ''.join(sentences[start:b+1]).strip()
            if segment_text:
                segments.append(segment_text)
            else:
                segments.append('')
            start = b + 1
        last_text = ''.join(sentences[start:]).strip()
        segments.append(last_text)

        # 如因边界选择不充分导致段数不足或过多，做轻微修正
        if len(segments) < num_segments:
            # 从尾段开始字符级补切，直到达到 num_segments
            while len(segments) < num_segments:
                last = segments.pop() if segments else ''
                if len(last) <= 1:
                    segments.extend([last, ''])
                else:
                    mid = len(last) // 2
                    segments.extend([last[:mid], last[mid:]])
        elif len(segments) > num_segments:
            # 合并最短相邻两段
            while len(segments) > num_segments and len(segments) >= 2:
                # 找到最短相邻对
                min_i = 0
                min_sum = float('inf')
                for i2 in range(len(segments) - 1):
                    ssum = len(segments[i2]) + len(segments[i2+1])
                    if ssum < min_sum:
                        min_sum = ssum
                        min_i = i2
                merged = segments[min_i] + segments[min_i+1]
                segments = segments[:min_i] + [merged] + segments[min_i+2:]
        return segments[:num_segments]

    # 3) 若句子数量 < 段数：字符级等分，尽量均衡
    # 等分为 num_segments 段
    base = total_len // num_segments
    rem = total_len % num_segments
    result: List[str] = []
    start_idx = 0
    for i in range(num_segments):
        length = base + (1 if i < rem else 0)
        end_idx = start_idx + length
        result.append(text[start_idx:end_idx])
        start_idx = end_idx
    return result

################ Keywords Extraction ################
def extract_keywords(server: str, model: str, script_data: Dict[str, Any]) -> Dict[str, Any]:
    """
    关键词提取 - 第二次LLM处理
    为每个段落提取关键词和氛围词
    """
    try:
        segments_text = []
        for segment in script_data["segments"]:
            segments_text.append(f"第{segment['index']}段: {segment['content']}")
        
        user_message = f"""请为以下每个段落提取关键词和氛围词，用于图像生成：

{chr(10).join(segments_text)}
"""
        
        output = text_to_text(
            server=server,
            model=model,
            prompt=user_message,
            system_message=keywords_extraction_prompt,
            max_tokens=4096,
            temperature=config.LLM_TEMPERATURE_KEYWORDS
        )
        
        if output is None:
            raise ValueError("未能从 API 获取响应。")
        
        # 鲁棒解析（先常规，失败则修复）
        keywords_data = parse_json_robust(output)
        
        # 精简对齐：按脚本段数对齐（多截断、少补空），不再额外校验/告警
        expected = len(script_data["segments"])  # 以脚本段数为准
        segs = list(keywords_data.get("segments") or [])
        keywords_data["segments"] = (
            segs[:expected]
            + [{"keywords": [], "atmosphere": []}] * max(0, expected - len(segs))
        )
        
        # 添加模型信息
        keywords_data["model_info"] = {
            "llm_server": server,
            "llm_model": model,
            "generation_type": "keywords_extraction"
        }
        keywords_data["created_time"] = datetime.datetime.now().isoformat()
        
        return keywords_data
    
    except json.JSONDecodeError:
        raise ValueError("解析关键词 JSON 输出失败")
    except Exception as e:
        raise ValueError(f"关键词提取错误: {e}")

################ Image Generation ################
def generate_opening_image(model: str, opening_style: str, 
                           image_size: str, output_dir: str) -> Optional[str]:
    """
    生成开场图像，使用预设风格。
    
    Args:
        model: 图像生成模型  
        opening_style: 开场图像风格键名 (minimal, tech, nature, abstract, vintage)
        image_size: 图像尺寸
        output_dir: 输出目录
    
    Returns:
        Optional[str]: 生成的开场图像路径，失败时返回None
    """
    try:
        # 获取预设风格提示词
        prompt = OPENING_IMAGE_STYLES.get(opening_style)
        if not prompt:
            # 使用第一个风格作为默认值
            default_style = next(iter(OPENING_IMAGE_STYLES))
            logger.warning(f"未找到开场图像风格: {opening_style}，使用默认风格: {default_style}")
            prompt = OPENING_IMAGE_STYLES[default_style]
        prompt = str(prompt).strip()
        
        # 调用豆包图像生成API
        image_url = text_to_image_doubao(
            prompt=prompt,
            size=image_size,
            model=model
        )

        if not image_url:
            raise ValueError("开场图像生成失败")

        # 下载并保存
        ensure_directory_exists(output_dir)
        image_path = os.path.join(output_dir, "opening.png")
        download_to_path(image_url, image_path, error_msg="开场图像下载失败")
        print(f"开场图像已保存: {image_path} (风格: {opening_style})")
        return image_path
    except Exception as e:
        logger.warning(f"开场图像生成失败: {e}")
        return None
def generate_images_for_segments(server: str, model: str, keywords_data: Dict[str, Any], 
                                image_style_preset: str, image_size: str, output_dir: str) -> Dict[str, Any]:
    """
    为每个段落生成图像
    
    Args:
        server: 图像生成服务商
        model: 图像生成模型
        keywords_data: 关键词数据
        image_style_preset: 图像风格预设名称
        image_size: 图像尺寸
        output_dir: 输出目录
    
    Returns:
        Dict[str, Any]: 包含图像路径列表和失败记录的字典
            - "image_paths": List[str] - 成功生成的图像路径
            - "failed_segments": List[int] - 生成失败的段落编号
    """
    try:
        image_paths = []
        failed_segments = []
        
        # 获取图像风格字符串
        image_style = get_image_style(image_style_preset)
        logger.info(f"使用图像风格: {image_style_preset} -> {image_style}")
        
        for i, segment_keywords in enumerate(keywords_data["segments"], 1):
            keywords = segment_keywords.get("keywords", [])
            atmosphere = segment_keywords.get("atmosphere", [])
            
            # 构建图像提示词
            style_part = f"[风格] {image_style}" if image_style else ""
            content_parts = []
            content_parts.extend(keywords)
            content_parts.extend(atmosphere)
            content_part = f"[内容] {' | '.join(content_parts)}" if content_parts else ""
            
            # 用换行符分隔不同部分，提高可读性
            prompt_sections = [part for part in [style_part, content_part] if part]
            final_prompt = "\n".join(prompt_sections)
            
            print(f"正在生成第{i}段图像...")
            
            # 重试逻辑：最多尝试3次
            success = False
            for attempt in range(3):
                try:
                    # 调用豆包图像生成API
                    image_url = text_to_image_doubao(
                        prompt=final_prompt,
                        size=image_size,
                        model=model
                    )
                    
                    if image_url:
                        # 下载并保存图像
                        image_path = os.path.join(output_dir, f"segment_{i}.png")
                        download_to_path(image_url, image_path, error_msg=f"下载第{i}段图像失败")
                        image_paths.append(image_path)
                        print(f"第{i}段图像已保存: {image_path}")
                        success = True
                        break
                    else:
                        if attempt < 2:  # 不是最后一次尝试
                            print(f"⚠️  第{i}段图像生成失败，正在重试 ({attempt + 2}/3)...")
                            continue
                        
                except Exception as e:
                    error_msg = str(e)
                    # 检查是否是敏感词错误
                    is_sensitive_error = (
                        "OutputImageSensitiveContentDetected" in error_msg or 
                        "sensitive" in error_msg.lower() or
                        "content" in error_msg.lower()
                    )
                    
                    if attempt < 2:  # 不是最后一次尝试
                        if is_sensitive_error:
                            print(f"⚠️  第{i}段图像涉及敏感内容，正在重试 ({attempt + 2}/3)...")
                        else:
                            print(f"⚠️  第{i}段图像生成失败: {error_msg}，正在重试 ({attempt + 2}/3)...")
                        continue
                    else:
                        # 最后一次尝试也失败了
                        if is_sensitive_error:
                            print(f"❌ 第{i}段图像涉及敏感内容，已跳过")
                        else:
                            print(f"❌ 第{i}段图像生成失败: {error_msg}，已跳过")
            
            # 如果3次尝试都失败，记录失败的段落
            if not success:
                failed_segments.append(i)
                # 添加空字符串占位，保持索引对应关系
                image_paths.append("")
        
        return {
            "image_paths": image_paths,
            "failed_segments": failed_segments
        }
    
    except Exception as e:
        raise ValueError(f"图像生成错误: {e}")

################ Voice Synthesis ################
def synthesize_voice_for_segments(server: str, voice: str, script_data: Dict[str, Any], output_dir: str) -> List[str]:
    """
    为每个段落合成语音
    """
    try:
        audio_paths = []
        
        for segment in script_data["segments"]:
            segment_index = segment["index"]
            content = segment["content"]
            
            print(f"正在生成第{segment_index}段语音...")
            
            # 生成语音文件路径：voice_{序号}.wav
            audio_filename = f"voice_{segment_index}.wav"
            audio_path = os.path.join(output_dir, audio_filename)
            
            # 调用语音合成API - 根据语音音色智能选择接口
            if server == "bytedance":
                # 使用字节语音合成大模型接口
                success = text_to_audio_bytedance(
                    text=content,
                    output_filename=audio_path,
                    voice=voice
                )
            else:
                raise ValueError(f"不支持的TTS服务商: {server}")
            
            if success:
                audio_paths.append(audio_path)
                print(f"第{segment_index}段语音已保存: {audio_path}")
            else:
                raise ValueError(f"生成第{segment_index}段语音失败")
        
        return audio_paths
    
    except Exception as e:
        raise ValueError(f"语音合成错误: {e}")

################ Video Composition ################
def compose_final_video(image_paths: List[str], audio_paths: List[str], output_path: str, 
                       script_data: Dict[str, Any] = None, enable_subtitles: bool = False,
                       bgm_audio_path: Optional[str] = None, bgm_volume: float = 0.15,
                       narration_volume: float = 1.0,
                       opening_image_path: Optional[str] = None,
                       opening_golden_quote: Optional[str] = None,
                       opening_narration_audio_path: Optional[str] = None) -> str:
    """
    合成最终视频
    
    Args:
        image_paths: 图像文件路径列表
        audio_paths: 音频文件路径列表
        output_path: 输出视频路径
        script_data: 脚本数据，用于生成字幕
        enable_subtitles: 是否启用字幕
    
    Returns:
        str: 输出视频路径
    """
    try:
        if len(image_paths) != len(audio_paths):
            raise ValueError("图像文件数量与音频文件数量不匹配")
        
        video_clips = []
        audio_clips = []
        
        # 可选：创建开场片段（图像 + 居中金句 + 可选开场口播）
        # 开场时长逻辑：若提供开场口播 => 时长=口播时长+OPENING_HOLD_AFTER_NARRATION_SECONDS；否则无开场
        opening_seconds = 0.0
        opening_voice_clip = None
        # 若提供开场口播音频，则以“音频长度 + 停留时长”作为总开场时长
        try:
            if opening_narration_audio_path and os.path.exists(opening_narration_audio_path):
                opening_voice_clip = AudioFileClip(opening_narration_audio_path)
                hold_after = float(getattr(config, "OPENING_HOLD_AFTER_NARRATION_SECONDS", 2.0))
                opening_seconds = float(opening_voice_clip.duration) + max(0.0, hold_after)
        except Exception as _oaerr:
            logger.warning(f"开场口播音频加载失败: {_oaerr}，将退回固定时长开场")
            opening_voice_clip = None
        
        if opening_image_path and os.path.exists(opening_image_path) and opening_seconds > 1e-3:
            try:
                print("正在创建开场片段…")
                opening_base = ImageClip(opening_image_path).with_duration(opening_seconds)

                # 解析可用字体（参考字幕配置）
                subtitle_config = config.SUBTITLE_CONFIG.copy()

                resolved_font = resolve_font_path(subtitle_config.get("font_family"))
                quote_text = (opening_golden_quote or "").strip()
                if quote_text:
                    # 读取开场金句样式（带默认值回退）
                    quote_style = getattr(config, "OPENING_QUOTE_STYLE", {}) or {}
                    base_font = int(config.SUBTITLE_CONFIG.get("font_size", 36))
                    scale = float(quote_style.get("font_scale", 1.3))
                    font_size = int(quote_style.get("font_size", base_font * scale))
                    text_color = quote_style.get("color", config.SUBTITLE_CONFIG.get("color", "white"))
                    stroke_color = quote_style.get("stroke_color", config.SUBTITLE_CONFIG.get("stroke_color", "black"))
                    stroke_width = int(quote_style.get("stroke_width", max(3, int(config.SUBTITLE_CONFIG.get("stroke_width", 3)))))
                    pos = quote_style.get("position", ("center", "center"))

                    # 开场金句换行：按 max_chars_per_line 和 max_lines 控制
                    try:
                        max_chars = int(quote_style.get("max_chars_per_line", 18))
                        max_q_lines = int(quote_style.get("max_lines", 4))
                        # 复用字幕拆分逻辑，严格按每行字符数限制
                        candidate_lines = split_text_for_subtitle(quote_text, max_chars, max_q_lines)
                        wrapped_quote = "\n".join(candidate_lines[:max_q_lines]) if candidate_lines else quote_text
                    except Exception:
                        wrapped_quote = quote_text

                    # 覆盖字体解析（优先采用 OPENING_QUOTE_STYLE.font_family）
                    font_override = quote_style.get("font_family")
                    if font_override and os.path.exists(font_override):
                        resolved_font = font_override

                    # 行间距与字间距（MoviePy 2.x 无直接参数，这里通过逐行排版+空格近似实现）
                    line_spacing_px = int(quote_style.get("line_spacing", 0))
                    letter_spaces = int(quote_style.get("letter_spacing", 0))

                    def _apply_letter_spacing(s: str, n: int) -> str:
                        if n <= 0 or not s:
                            return s
                        return (" " * n).join(list(s))

                    def _make_text_clip(text: str) -> TextClip:
                        return TextClip(
                            text=text,
                            font_size=font_size,
                            color=text_color,
                            font=resolved_font or config.SUBTITLE_CONFIG.get("font_family"),
                            stroke_color=stroke_color,
                            stroke_width=stroke_width
                        )

                    try:
                        # 预处理字间距
                        lines = wrapped_quote.split("\n") if wrapped_quote else []
                        lines = [_apply_letter_spacing(ln, letter_spaces) for ln in lines] if lines else []

                        # 无需自定义行距：直接用单 TextClip（多行通过 \n 渲染）
                        if line_spacing_px <= 0 or not (isinstance(pos, tuple) and pos == ("center", "center")):
                            processed = "\n".join(lines) if lines else wrapped_quote
                            text_clip = _make_text_clip(processed).with_position(pos).with_duration(opening_seconds)
                            opening_clip = CompositeVideoClip([opening_base, text_clip])
                        else:
                            # 居中且需要行距：逐行排布
                            video_w, video_h = opening_base.size
                            line_clips: List[Any] = [_make_text_clip(ln) for ln in lines] if lines else []
                            if line_clips:
                                total_h = sum(c.h for c in line_clips) + line_spacing_px * (len(line_clips) - 1)
                                y_start = max(0, (video_h - total_h) // 2)
                                y_cur = y_start
                                placed: List[Any] = [opening_base]
                                for c in line_clips:
                                    placed.append(c.with_position(("center", y_cur)).with_duration(opening_seconds))
                                    y_cur += c.h + line_spacing_px
                                opening_clip = CompositeVideoClip(placed)
                            else:
                                text_clip = _make_text_clip(wrapped_quote).with_position(pos).with_duration(opening_seconds)
                                opening_clip = CompositeVideoClip([opening_base, text_clip])
                    except Exception:
                        text_clip = _make_text_clip(wrapped_quote).with_position(pos).with_duration(opening_seconds)
                        opening_clip = CompositeVideoClip([opening_base, text_clip])
                else:
                    opening_clip = opening_base

                # 绑定开场口播音频（如存在）
                if opening_voice_clip is not None:
                    try:
                        opening_clip = opening_clip.with_audio(opening_voice_clip)
                    except Exception as _bindaerr:
                        logger.warning(f"为开场片段绑定音频失败: {_bindaerr}")

                # 添加开场片段渐隐效果：在口播结束后的停留时间内逐渐变黑
                try:
                    if opening_voice_clip is not None and hold_after > 1e-3:
                        voice_duration = float(opening_voice_clip.duration)
                        fade_start_time = voice_duration  # 口播结束时开始渐隐
                        fade_duration = hold_after        # 渐隐持续时间
                        
                        def _opening_fade_out(gf, t):
                            try:
                                if t < fade_start_time:
                                    # 口播期间：正常显示
                                    return gf(t)
                                elif t >= opening_seconds:
                                    # 超出总时长：完全黑屏
                                    return 0.0 * gf(t)
                                else:
                                    # 渐隐期间：线性递减alpha
                                    fade_progress = (t - fade_start_time) / fade_duration
                                    alpha = max(0.0, 1.0 - fade_progress)
                                    return alpha * gf(t)
                            except Exception:
                                return gf(t)
                        
                        opening_clip = opening_clip.transform(_opening_fade_out, keep_duration=True)
                        print(f"🎬 已为开场片段添加{hold_after}s渐隐效果")
                except Exception as fade_err:
                    logger.warning(f"开场片段渐隐效果添加失败: {fade_err}")

                video_clips.append(opening_clip)
            except Exception as e:
                logger.warning(f"开场片段生成失败: {e}，将跳过开场")

        # 为每个段落创建视频片段
        for i, (image_path, audio_path) in enumerate(zip(image_paths, audio_paths)):
            print(f"正在处理第{i+1}段视频...")
            
            # 加载音频获取时长
            audio_clip = AudioFileClip(audio_path)
            duration = audio_clip.duration
            
            # 创建图像剪辑，设置持续时间为音频长度 (MoviePy 2.x 使用 with_duration)
            image_clip = ImageClip(image_path).with_duration(duration)
            
            # 组合图像和音频 (MoviePy 2.x 使用 with_audio)
            video_clip = image_clip.with_audio(audio_clip)
            video_clips.append(video_clip)
            audio_clips.append(audio_clip)
        
        # 连接所有视频片段
        print("正在合成最终视频...")
        # 使用 compose 方式合并，避免音频轨丢失或不同尺寸导致的问题
        final_video = concatenate_videoclips(video_clips, method="compose")
        
        # 添加字幕（如果启用）
        # 生效的字幕开关需同时满足：运行时参数与全局配置均为 True
        effective_subtitles = bool(enable_subtitles) and bool(getattr(config, "SUBTITLE_CONFIG", {}).get("enabled", True))
        if effective_subtitles and script_data:
            print("正在添加字幕...")
            try:
                # 传入最终视频尺寸，便于字幕计算边距/背景
                subtitle_config = config.SUBTITLE_CONFIG.copy()
                subtitle_config["video_size"] = final_video.size
                # 传入每段音频真实时长用于精准对齐
                subtitle_config["segment_durations"] = [ac.duration for ac in audio_clips]
                # 开场字幕偏移：让第一段字幕从开场片段之后开始
                subtitle_config["offset_seconds"] = opening_seconds
                subtitle_clips = create_subtitle_clips(script_data, subtitle_config)
                if subtitle_clips:
                    # 将字幕与视频合成
                    final_video = CompositeVideoClip([final_video] + subtitle_clips)
                    print(f"已添加 {len(subtitle_clips)} 个字幕剪辑")
                else:
                    print("未生成任何字幕剪辑")
            except Exception as e:
                logger.warning(f"添加字幕失败: {str(e)}，继续生成无字幕视频")

        # 调整口播音量（在与BGM混音前）——MoviePy 2.x 使用 MultiplyVolume
        try:
            if final_video.audio is not None and narration_volume is not None:
                narration_audio = final_video.audio
                if isinstance(narration_volume, (int, float)) and abs(float(narration_volume) - 1.0) > 1e-9:
                    narration_audio = narration_audio.with_effects([MultiplyVolume(float(narration_volume))])
                    final_video = final_video.with_audio(narration_audio)
                    print(f"🔊 口播音量调整为: {float(narration_volume)}")
        except Exception as e:
            logger.warning(f"口播音量调整失败: {str(e)}，将使用原始音量")

        # 在视频开头应用视觉渐显（从黑到正常）
        try:
            fade_in_seconds = float(getattr(config, "OPENING_FADEIN_SECONDS", 0.0))
            if fade_in_seconds > 1e-3:
                def _fade_in_frame(gf, t):
                    try:
                        alpha = min(1.0, max(0.0, float(t) / float(fade_in_seconds)))
                    except Exception:
                        alpha = 1.0
                    return alpha * gf(t)
                try:
                    final_video = final_video.transform(_fade_in_frame, keep_duration=True)
                    print(f"🎬 已添加开场渐显 {fade_in_seconds}s")
                except Exception as _ferr:
                    logger.warning(f"开场渐显应用失败: {_ferr}")
        except Exception as e:
            logger.warning(f"读取开场渐显配置失败: {e}")
        
        # 在片尾追加 config.ENDING_FADE_SECONDS 秒静帧并渐隐（仅画面，无口播音频）
        try:
            tail_seconds = float(getattr(config, "ENDING_FADE_SECONDS", 2.5))
            if isinstance(image_paths, list) and len(image_paths) > 0 and tail_seconds > 1e-3:
                last_image_path = image_paths[-1]
                tail_clip = ImageClip(last_image_path).with_duration(tail_seconds)
                # 使用 transform 实现到黑场的线性渐隐
                def _fade_frame(gf, t):
                    try:
                        alpha = max(0.0, 1.0 - float(t) / float(tail_seconds))
                    except Exception:
                        alpha = 0.0
                    return alpha * gf(t)
                try:
                    tail_clip = tail_clip.transform(_fade_frame, keep_duration=True)
                except Exception:
                    pass
                final_video = concatenate_videoclips([final_video, tail_clip], method="compose")
                print(f"🎬 已添加片尾静帧 {tail_seconds}s 并渐隐")
        except Exception as tail_err:
            logger.warning(f"片尾静帧添加失败: {tail_err}，将继续生成无片尾渐隐的视频")
        
        # 可选：叠加背景音乐（与口播混音）
        bgm_clip = None
        try:
            if bgm_audio_path and os.path.exists(bgm_audio_path):
                print(f"🎵 开始处理背景音乐: {bgm_audio_path}")
                bgm_clip = AudioFileClip(bgm_audio_path)
                print(f"🎵 BGM加载成功，时长: {bgm_clip.duration:.2f}秒")
                
                # 调整 BGM 音量（MoviePy 2.x MultiplyVolume）
                try:
                    if isinstance(bgm_volume, (int, float)) and abs(float(bgm_volume) - 1.0) > 1e-9:
                        bgm_clip = bgm_clip.with_effects([MultiplyVolume(float(bgm_volume))])
                        print(f"🎵 BGM音量调整为: {float(bgm_volume)}")
                except Exception:
                    print("⚠️ BGM音量调整失败，使用原音量")
                    pass
                
                # 循环或裁剪至视频总时长（优先使用 MoviePy 2.x 的 AudioLoop）
                try:
                    target_duration = final_video.duration
                    print(f"🎵 视频总时长: {target_duration:.2f}秒，BGM时长: {bgm_clip.duration:.2f}秒")
                    if AudioLoop is not None:
                        # 使用 2.x 的 AudioLoop 效果类
                        bgm_clip = bgm_clip.with_effects([AudioLoop(duration=target_duration)])
                        print(f"🎵 BGM长度适配完成（AudioLoop），最终时长: {bgm_clip.duration:.2f}秒")
                    else:
                        # 简化的回退：直接裁剪到目标时长（避免复杂手动循环）
                        if hasattr(bgm_clip, "with_duration"):
                            bgm_clip = bgm_clip.with_duration(min(bgm_clip.duration, target_duration))
                            print("⚠️ AudioLoop 不可用，已将BGM裁剪到目标时长")
                        else:
                            raise RuntimeError("AudioLoop 不可用，且不支持 with_duration")

                except Exception as loop_err:
                    print(f"⚠️ 背景音乐长度适配失败: {loop_err}，将不添加BGM继续生成")
                    logger.warning(f"背景音乐循环/裁剪失败: {loop_err}，将不添加BGM继续生成")
                    bgm_clip = None
                    
                # 合成复合音频
                if bgm_clip is not None:
                    print("🎵 开始合成背景音乐和口播音频")
                    # 通用线性淡出增益函数（用于片尾淡出）
                    def _linear_fade_out_gain(total: float, tail: float):
                        cutoff = max(0.0, total - tail)
                        def _gain_any(t_any):
                            import numpy as _np
                            def _scalar(ts: float) -> float:
                                if ts <= cutoff:
                                    return 1.0
                                if ts >= total:
                                    return 0.0
                                return max(0.0, 1.0 - (ts - cutoff) / tail)
                            if hasattr(t_any, "__len__"):
                                return _np.array([_scalar(float(ts)) for ts in t_any])
                            return _scalar(float(t_any))
                        return _gain_any
                    if final_video.audio is not None:
                        # 可选：自动 Ducking，根据口播包络动态压低 BGM（MoviePy 2.x 通过 transform 实现时间变增益）
                        try:
                            if getattr(config, "AUDIO_DUCKING_ENABLED", False):
                                strength = float(getattr(config, "AUDIO_DUCKING_STRENGTH", 0.7))
                                smooth_sec = float(getattr(config, "AUDIO_DUCKING_SMOOTH_SECONDS", 0.12))
                                total_dur = float(final_video.duration)
                                # 采样频率（包络计算），20Hz 足够平滑且开销低
                                env_fps = 20.0
                                num_samples = max(2, int(total_dur * env_fps) + 1)
                                times = np.linspace(0.0, total_dur, num_samples)
                                # 估算口播瞬时幅度（绝对值，通道取均值）
                                amp = np.zeros_like(times)
                                for i, t in enumerate(times):
                                    try:
                                        frame = final_video.audio.get_frame(float(min(max(0.0, t), total_dur - 1e-6)))
                                        # frame 形如 [L, R]
                                        amp[i] = float(np.mean(np.abs(frame)))
                                    except Exception:
                                        amp[i] = 0.0
                                # 平滑（简单滑动平均窗口）
                                win = max(1, int(smooth_sec * env_fps))
                                if win > 1:
                                    kernel = np.ones(win, dtype=float) / win
                                    amp = np.convolve(amp, kernel, mode="same")
                                # 归一化
                                max_amp = float(np.max(amp)) if np.max(amp) > 1e-8 else 1.0
                                env = amp / max_amp
                                # 计算 duck 增益曲线：口播强 -> BGM 更低
                                gains = 1.0 - strength * env
                                gains = np.clip(gains, 0.0, 1.0)
                                # 构建时间变增益函数（支持标量/向量 t）
                                def _gain_lookup(t_any):
                                    import numpy as _np
                                    def _lookup_scalar(ts: float) -> float:
                                        if ts <= 0.0:
                                            return float(gains[0])
                                        if ts >= total_dur:
                                            return float(gains[-1])
                                        idx = int(ts * env_fps)
                                        if idx < 0:
                                            idx = 0
                                        if idx >= gains.shape[0]:
                                            idx = gains.shape[0] - 1
                                        return float(gains[idx])
                                    if hasattr(t_any, "__len__"):
                                        return _np.array([_lookup_scalar(float(ts)) for ts in t_any])
                                    return _lookup_scalar(float(t_any))

                                # 应用时间变增益到 BGM（使用 transform），注意多声道广播维度
                                bgm_clip = bgm_clip.transform(
                                    lambda gf, t: (
                                        (_gain_lookup(t)[:, None] if hasattr(t, "__len__") else _gain_lookup(t))
                                        * gf(t)
                                    ),
                                    keep_duration=True,
                                )
                                print(f"🎚️ 已启用自动Ducking（strength={strength}, smooth={smooth_sec}s）")
                        except Exception as duck_err:
                            logger.warning(f"自动Ducking失败: {duck_err}，将使用恒定音量BGM")
                        # 在片尾对 BGM 做淡出（不影响口播，因为尾段无口播）
                        try:
                            total_dur = float(final_video.duration)
                            fade_tail = float(getattr(config, "ENDING_FADE_SECONDS", 2.5))
                            fade_gain = _linear_fade_out_gain(total_dur, fade_tail)
                            bgm_clip = bgm_clip.transform(
                                lambda gf, t: ((fade_gain(t)[:, None]) if hasattr(t, "__len__") else fade_gain(t)) * gf(t),
                                keep_duration=True,
                            )
                            print(f"🎚️ 已添加BGM片尾{fade_tail}s淡出")
                        except Exception as _fade_err:
                            logger.warning(f"BGM淡出应用失败: {_fade_err}")
                        mixed_audio = CompositeAudioClip([final_video.audio, bgm_clip])
                        print("🎵 BGM与口播音频合成完成")
                    else:
                        # 无口播，仅 BGM；同样添加片尾淡出
                        try:
                            total_dur = float(final_video.duration)
                            fade_tail = float(getattr(config, "ENDING_FADE_SECONDS", 2.5))
                            fade_gain = _linear_fade_out_gain(total_dur, fade_tail)
                            bgm_clip = bgm_clip.transform(
                                lambda gf, t: ((fade_gain(t)[:, None]) if hasattr(t, "__len__") else fade_gain(t)) * gf(t),
                                keep_duration=True,
                            )
                            print(f"🎚️ 已添加BGM片尾{fade_tail}s淡出")
                        except Exception as _fade_err:
                            logger.warning(f"BGM淡出应用失败: {_fade_err}")
                        mixed_audio = CompositeAudioClip([bgm_clip])
                        print("🎵 仅添加BGM音频（无口播音频）")
                    final_video = final_video.with_audio(mixed_audio)
                    print("🎵 背景音乐添加成功！")
                else:
                    print("❌ BGM处理失败，生成无背景音乐视频")
            else:
                if bgm_audio_path:
                    print(f"⚠️ 背景音乐文件不存在: {bgm_audio_path}")
                else:
                    print("ℹ️ 未指定背景音乐文件")
        except Exception as e:
            print(f"❌ 背景音乐处理异常: {str(e)}")
            logger.warning(f"背景音乐处理失败: {str(e)}，将继续生成无背景音乐的视频")

        # 输出最终视频：使用简单进度条，避免某些终端环境下 tqdm 多行滚动刷屏
        moviepy_logger = 'bar'

        try:
            # 优先尝试 macOS 硬件编码（如 ffmpeg 启用了 videotoolbox，将显著加速）
            final_video.write_videofile(
                output_path,
                fps=15,
                codec='h264_videotoolbox',
                audio_codec='aac',
                bitrate='5M',
                ffmpeg_params=['-pix_fmt', 'yuv420p', '-movflags', '+faststart'],
                logger=moviepy_logger
            )
        except Exception as _hw_err:
            print(f"⚠️ 硬件编码不可用或失败，回退到软件编码: {_hw_err}")
            # 回退到软件 x264：保持 CRF 画质一致，同时提升 preset 和线程以加速
            final_video.write_videofile(
                output_path,
                fps=15,
                codec='libx264',
                audio_codec='aac',
                preset='veryfast',
                threads=os.cpu_count() or 4,
                ffmpeg_params=['-crf', '23', '-pix_fmt', 'yuv420p', '-movflags', '+faststart'],
                logger=moviepy_logger
            )
        
        # 释放资源
        for clip in video_clips:
            clip.close()
        for aclip in audio_clips:
            aclip.close()
        final_video.close()
        if bgm_clip is not None:
            try:
                bgm_clip.close()
            except Exception:
                pass
        if 'opening_voice_clip' in locals() and opening_voice_clip is not None:
            try:
                opening_voice_clip.close()
            except Exception:
                pass
        
        print(f"最终视频已保存: {output_path}")
        return output_path
    
    except Exception as e:
        raise ValueError(f"视频合成错误: {e}")

################ Style Helper Functions ################
def get_image_style(style_name: str = "cinematic") -> str:
    """
    获取图像风格字符串
    
    Args:
        style_name: 风格名称，如果不存在则返回第一个风格
    
    Returns:
        str: 图像风格描述字符串
    """
    return IMAGE_STYLE_PRESETS.get(style_name, list(IMAGE_STYLE_PRESETS.values())[0])

def _split_text_evenly(text: str, max_chars_per_line: int) -> List[str]:
    """
    将文本均匀切分，避免出现过短的尾段
    
    Args:
        text: 要切分的文本
        max_chars_per_line: 每行最大字符数
    
    Returns:
        List[str]: 均匀切分后的文本片段列表
    """
    if len(text) <= max_chars_per_line:
        return [text]
    
    # 计算需要几段以及每段的理想长度
    total_chars = len(text)
    num_segments = (total_chars + max_chars_per_line - 1) // max_chars_per_line  # 向上取整
    
    # 计算每段的理想长度（尽可能均匀）
    base_length = total_chars // num_segments
    remainder = total_chars % num_segments
    
    # 分配长度：前remainder段多分配1个字符
    result = []
    start = 0
    
    for i in range(num_segments):
        # 前remainder段长度为base_length+1，后面的段长度为base_length
        length = base_length + (1 if i < remainder else 0)
        end = start + length
        result.append(text[start:end])
        start = end
    
    return result

def split_text_for_subtitle(text: str, max_chars_per_line: int = 20, max_lines: int = 2) -> List[str]:
    """
    将长文本分割为适合字幕显示的短句，采用分层切分策略：
    1. 第一层：按句号等重标点切分
    2. 第二层：对超长片段按逗号等轻标点切分  
    3. 第三层：对仍超长的片段按字符数均匀硬切
    """
    def _split_by_light_punctuation(fragment: str) -> List[str]:
        """第二层：智能配对标点切分，保持引用内容完整"""
        # 配对标点映射（含中文引号、书名号、日文引号）
        paired_punctuation = {
            '"': '"',
            "'": "'",
            '《': '》',
            '“': '”',
            '‘': '’',
            '「': '」',
        }
        
        # 第一步：提取配对标点内容和普通文本
        segments = []
        i = 0
        
        while i < len(fragment):
            # 查找下一个配对标点的开始
            next_pair_start = float('inf')
            pair_type = None
            
            for start_p in paired_punctuation:
                pos = fragment.find(start_p, i)
                if pos != -1 and pos < next_pair_start:
                    next_pair_start = pos
                    pair_type = start_p
            
            if next_pair_start == float('inf'):
                # 没有更多配对标点，剩余文本作为普通段
                if i < len(fragment):
                    segments.append(('text', fragment[i:]))
                break
            
            # 添加配对标点前的普通文本
            if next_pair_start > i:
                segments.append(('text', fragment[i:next_pair_start]))
            
            # 查找配对的结束标点
            end_punct = paired_punctuation[pair_type]
            end_pos = fragment.find(end_punct, next_pair_start + 1)
            
            if end_pos != -1:
                # 找到配对，提取完整配对内容
                paired_content = fragment[next_pair_start:end_pos + 1]
                segments.append(('paired', paired_content))
                i = end_pos + 1
            else:
                # 没找到配对，当作普通字符处理
                segments.append(('text', fragment[next_pair_start:next_pair_start + 1]))
                i = next_pair_start + 1
        
        # 第二步：处理每个段落
        final_parts = []
        
        for seg_type, content in segments:
            if seg_type == 'paired':
                # 配对内容保持完整（但检查是否超长）
                if len(content) <= max_chars_per_line:
                    final_parts.append(content)
                else:
                    # 配对内容太长，需要硬切（保持配对标点完整性的前提下）
                    final_parts.extend(_split_text_evenly(content, max_chars_per_line))
            else:
                # 普通文本应用轻标点切分（不包含配对标点）
                if len(content) <= max_chars_per_line:
                    final_parts.append(content)
                else:
                    # 对普通文本应用简化的轻标点切分
                    text_parts = _split_simple_punctuation(content)
                    final_parts.extend(text_parts)
        
        return [part for part in final_parts if part.strip()]
    
    def _split_simple_punctuation(text: str) -> List[str]:
        """对普通文本进行简化的轻标点切分（不处理配对标点）"""
        tokens = re.split(r'([，、,:："])', text)
        if len(tokens) <= 1:
            return _split_text_evenly(text, max_chars_per_line)
        
        parts = []
        buf = ""
        for token in tokens:
            if not token:
                continue
            if len(buf + token) <= max_chars_per_line:
                buf += token
            else:
                if buf:
                    parts.append(buf)
                buf = token
        if buf:
            parts.append(buf)
        
        # 检查是否还有超长片段需要硬切
        final_parts = []
        for part in parts:
            if len(part) <= max_chars_per_line:
                final_parts.append(part)
            else:
                final_parts.extend(_split_text_evenly(part, max_chars_per_line))
        
        return final_parts
    
    # 如果文本不超长，直接返回
    if len(text) <= max_chars_per_line:
        return [text]
    
    # 第一层：按句号等重标点分割
    raw_parts = re.split(r'([。！？；.!?])', text)
    
    # 重新组合：将标点符号附加到前面的文本
    sentences = []
    i = 0
    while i < len(raw_parts):
        if not raw_parts[i].strip():
            i += 1
            continue
        
        # 当前是文本内容
        current = raw_parts[i].strip()
        
        # 检查下一个是否是标点符号
        if i + 1 < len(raw_parts) and raw_parts[i + 1].strip() in '。！？；.!?':
            current += raw_parts[i + 1].strip()
            i += 2
        else:
            i += 1
        
        if current:
            sentences.append(current)
    
    # 如果没有重标点，跳到第二层处理
    if not sentences:
        sentences = [text]
    
    # 逐片段处理
    result = []
    
    for sentence in sentences:
        if len(sentence) <= max_chars_per_line:
            # 第一层切分结果符合要求，直接保留
            result.append(sentence)
        else:
            # 第一层切分结果超长，进入第二层
            second_level_parts = _split_by_light_punctuation(sentence)
            
            for part in second_level_parts:
                if len(part) <= max_chars_per_line:
                    # 第二层切分结果符合要求，直接保留
                    result.append(part)
                else:
                    # 第二层切分结果仍超长，进入第三层硬切
                    result.extend(_split_text_evenly(part, max_chars_per_line))
    
    # 注释：移除 max_lines 截断逻辑，确保所有分割的字幕都能轮播显示
    # 原逻辑会截断超出行数限制的字幕，导致只显示前几句
    
    return result

def create_subtitle_clips(script_data: Dict[str, Any], subtitle_config: Dict[str, Any] = None) -> List[TextClip]:
    """
    创建字幕剪辑列表
    
    Args:
        script_data: 脚本数据，包含segments信息
        subtitle_config: 字幕配置，如果为None则使用默认配置
    
    Returns:
        List[TextClip]: 字幕剪辑列表
    """
    if subtitle_config is None:
        subtitle_config = config.SUBTITLE_CONFIG.copy()
    
    subtitle_clips = []
    current_time = float(subtitle_config.get("offset_seconds", 0.0) if isinstance(subtitle_config, dict) else 0.0)
    
    logger.info("开始创建字幕剪辑...")
    
    # 解析可用字体（优先使用系统中的中文字体文件路径，避免中文缺字）
    def _resolve_font_path(preferred: Optional[str]) -> Optional[str]:
        return resolve_font_path(preferred)

    resolved_font = _resolve_font_path(subtitle_config.get("font_family"))
    if not resolved_font:
        logger.warning("未能解析到可用中文字体，可能导致字幕无法显示中文字符。建议在 config.SUBTITLE_CONFIG.font_family 指定字体文件路径。")

    # 读取视频尺寸（用于计算底部边距和背景条）
    video_size = subtitle_config.get("video_size", (1280, 720))
    video_width, video_height = video_size

    segment_durations = subtitle_config.get("segment_durations", [])

    # 定义需要替换为空格的标点集合（中英文常见标点）
    # 中文引号“”‘’与书名号《》保留；英文双引号直接替换为双空格；其他标点替换为双空格
    # 使用 + 将连续标点视作一个整体，避免产生过多空白
    punctuation_pattern = r"[-.,!?;:\"，。！？；：（）()\[\]{}【】—…–、]+"

    # 内部辅助：根据配置生成文本与可选阴影、背景条的剪辑列表
    def _make_text_and_bg_clips(display_text: str, start_time: float, duration: float) -> List[Any]:
        position = subtitle_config["position"]
        margin_bottom = int(subtitle_config.get("margin_bottom", 0))
        anchor_x = position[0] if isinstance(position, tuple) else "center"

        # 检查是否有多行文本，如果有则需要特殊处理行间距
        lines = display_text.split('\n') if '\n' in display_text else [display_text]
        line_spacing_px = int(subtitle_config.get("line_spacing", 15))
        
        clips_to_add: List[Any] = []
        
        if len(lines) > 1 and line_spacing_px > 0:
            # 多行文本：手动排版每一行，应用行间距
            line_clips = []
            for line in lines:
                if line.strip():  # 忽略空行
                    line_clip = TextClip(
                        text=line.strip(),
                        font_size=subtitle_config["font_size"],
                        color=subtitle_config["color"],
                        font=resolved_font or subtitle_config["font_family"],
                        stroke_color=subtitle_config["stroke_color"],
                        stroke_width=subtitle_config["stroke_width"]
                    )
                    line_clips.append(line_clip)
            
            if line_clips:
                # 计算总高度并居中
                total_height = sum(clip.h for clip in line_clips) + line_spacing_px * (len(line_clips) - 1)
                
                # 计算起始Y位置
                if isinstance(position, tuple) and len(position) == 2 and position[1] == "bottom":
                    baseline_safe_padding = int(subtitle_config.get("baseline_safe_padding", 4))
                    y_start = max(0, video_height - margin_bottom - total_height - baseline_safe_padding)
                elif isinstance(position, tuple) and len(position) == 2:
                    try:
                        y_start = max(0, (video_height - total_height) // 2) if position[1] == "center" else int(position[1])
                    except:
                        y_start = (video_height - total_height) // 2
                else:
                    y_start = (video_height - total_height) // 2
                
                # 放置每一行
                current_y = y_start
                for line_clip in line_clips:
                    main_pos = (anchor_x, current_y)
                    line_clip = line_clip.with_position(main_pos).with_start(start_time).with_duration(duration)
                    clips_to_add.append(line_clip)
                    
                    # 添加阴影（如果启用）
                    if subtitle_config.get("shadow_enabled", False):
                        shadow_color = subtitle_config.get("shadow_color", "black")
                        shadow_offset = subtitle_config.get("shadow_offset", (2, 2))
                        shadow_x = main_pos[0] if isinstance(main_pos[0], int) else 0
                        shadow_y = main_pos[1] if isinstance(main_pos[1], int) else current_y
                        try:
                            shadow_pos = (shadow_x + shadow_offset[0], shadow_y + shadow_offset[1])
                        except:
                            shadow_pos = main_pos
                            
                        shadow_clip = TextClip(
                            text=line_clip.text,
                            font_size=subtitle_config["font_size"],
                            color=shadow_color,
                            font=resolved_font or subtitle_config["font_family"]
                        ).with_position(shadow_pos).with_start(start_time).with_duration(duration)
                        clips_to_add.insert(-1, shadow_clip)  # 插入到文本前面
                    
                    current_y += line_clip.h + line_spacing_px
        else:
            # 单行文本：使用原来的逻辑
            main_clip = TextClip(
                text=display_text,
                font_size=subtitle_config["font_size"],
                color=subtitle_config["color"],
                font=resolved_font or subtitle_config["font_family"],
                stroke_color=subtitle_config["stroke_color"],
                stroke_width=subtitle_config["stroke_width"]
            )

            # 计算文本位置（当定位 bottom 时基于文本高度与边距）
            if isinstance(position, tuple) and len(position) == 2 and position[1] == "bottom":
                baseline_safe_padding = int(subtitle_config.get("baseline_safe_padding", 4))
                y_text = max(0, video_height - margin_bottom - main_clip.h - baseline_safe_padding)
                main_pos = (anchor_x, y_text)
            else:
                main_pos = position

            main_clip = main_clip.with_position(main_pos).with_start(start_time).with_duration(duration)

            # 可选阴影
            if subtitle_config.get("shadow_enabled", False):
                shadow_color = subtitle_config.get("shadow_color", "black")
                shadow_offset = subtitle_config.get("shadow_offset", (2, 2))
                shadow_x = main_pos[0] if isinstance(main_pos[0], int) else 0
                shadow_y = main_pos[1] if isinstance(main_pos[1], int) else 0
                try:
                    shadow_pos = (shadow_x + shadow_offset[0], shadow_y + shadow_offset[1])
                except:
                    shadow_pos = main_pos
                    
                shadow_clip = TextClip(
                    text=display_text,
                    font_size=subtitle_config["font_size"],
                    color=shadow_color,
                    font=resolved_font or subtitle_config["font_family"]
                ).with_position(shadow_pos).with_start(start_time).with_duration(duration)
                clips_to_add.extend([shadow_clip, main_clip])
            else:
                clips_to_add.append(main_clip)

        # 背景条（如启用）
        bg_color = subtitle_config.get("background_color")
        bg_opacity = float(subtitle_config.get("background_opacity", 0))
        if bg_color and bg_opacity > 0.0:
            bg_height = int(
                subtitle_config["font_size"] * subtitle_config.get("max_lines", 2)
                + subtitle_config.get("line_spacing", 10) + 4
            )
            bg_clip = ColorClip(size=(video_width, bg_height), color=bg_color)
            if hasattr(bg_clip, "with_opacity"):
                bg_clip = bg_clip.with_opacity(bg_opacity)
            y_bg = max(0, video_height - margin_bottom - bg_height)
            bg_clip = bg_clip.with_position(("center", y_bg)).with_start(start_time).with_duration(duration)
            clips_to_add.insert(0, bg_clip)  # 背景在最底层

        return clips_to_add

    for i, segment in enumerate(script_data["segments"], 1):
        content = segment["content"]
        # 优先使用真实音频时长，其次回退到估算时长
        duration = None
        if isinstance(segment_durations, list) and len(segment_durations) >= i:
            duration = float(segment_durations[i-1])
        if duration is None:
            duration = float(segment.get("estimated_duration", 0))
        
        logger.debug(f"处理第{i}段字幕，时长: {duration}秒")
        
        # 分割长文本为适合显示的字幕
        subtitle_texts = split_text_for_subtitle(
            content,
            subtitle_config["max_chars_per_line"],
            subtitle_config["max_lines"]
        )
        
        # 计算每行字幕的显示时长：按行字符数占比分配，确保总和==段时长
        subtitle_start_time = current_time
        line_durations: List[float] = []
        if len(subtitle_texts) > 0:
            lengths = [max(1, len(t)) for t in subtitle_texts]
            total_len = sum(lengths)
            acc = 0.0
            for idx, L in enumerate(lengths):
                if idx < len(lengths) - 1:
                    d = duration * (L / total_len)
                    line_durations.append(d)
                    acc += d
                else:
                    line_durations.append(max(0.0, duration - acc))
        else:
            line_durations = [duration]
        
        for subtitle_text, subtitle_duration in zip(subtitle_texts, line_durations):
            try:
                # 将连续标点替换为两个空格，并压缩可能产生的多余空格
                display_text = re.sub(punctuation_pattern, "  ", subtitle_text)
                display_text = re.sub(r" {3,}", "  ", display_text)

                clips_to_add = _make_text_and_bg_clips(display_text, subtitle_start_time, subtitle_duration)
                subtitle_clips.extend(clips_to_add)
                logger.debug(f"创建字幕: '{subtitle_text[:20]}...' 时间: {subtitle_start_time:.1f}-{subtitle_start_time+subtitle_duration:.1f}s")
                
                subtitle_start_time += subtitle_duration
                
            except Exception as e:
                logger.warning(f"创建字幕失败: {str(e)}，跳过此字幕")
                continue
        
        current_time += duration
    
    logger.info(f"字幕创建完成，共创建 {len(subtitle_clips)} 个字幕剪辑")
    return subtitle_clips