"""
文档处理器 - 处理文档格式转换
专门负责JSON数据与DOCX格式之间的转换

调用关系:
- core/pipeline.py: 调用export_script_to_docx导出阅读版DOCX文档
- 可被用户通过CLI或Web界面间接调用来进行文档格式转换
- 从utils.py迁移而来，专注于DOCX文档的导出和解析功能
"""

import os
from typing import Dict, Any

from core.utils import logger, ensure_directory_exists, FileProcessingError


def export_script_to_docx(script_data: Dict[str, Any], docx_path: str) -> str:
    """
    将脚本JSON导出为可阅读的DOCX文档
    
    Args:
        script_data: 含有 title 与 segments 的脚本数据
        docx_path: 输出的docx文件完整路径
        
    Returns:
        str: 实际保存的docx路径
    """
    from docx import Document
    from docx.enum.text import WD_ALIGN_PARAGRAPH

    document = Document()
    _setup_docx_style(document)

    title_text = script_data.get('title', 'untitled')
    segments = script_data.get('segments', []) or []

    # 标题（居中）
    title_para = document.add_paragraph()
    title_run = title_para.add_run(title_text)
    _setup_docx_run(title_run)
    _setup_docx_paragraph(title_para, WD_ALIGN_PARAGRAPH.CENTER)

    # 正文段落（两端对齐）
    for seg in segments:
        content = (seg or {}).get('content', '')
        if not content:
            continue
        p = document.add_paragraph()
        r = p.add_run(content)
        _setup_docx_run(r)
        _setup_docx_paragraph(p, WD_ALIGN_PARAGRAPH.JUSTIFY)

    ensure_directory_exists(os.path.dirname(docx_path))
    document.save(docx_path)
    logger.info(f"阅读版DOCX已保存: {docx_path}")
    return docx_path


def export_raw_to_docx(raw_data: Dict[str, Any], docx_path: str) -> str:
    """
    将原始LLM输出的JSON数据导出为带标记的DOCX文档，方便用户编辑
    
    Args:
        raw_data: 包含title、golden_quote、content的原始数据
        docx_path: 输出的docx文件路径
        
    Returns:
        str: 实际保存的docx路径
    """
    from docx import Document

    document = Document()
    _setup_docx_style(document)

    # 添加编辑说明
    instruction_para = document.add_paragraph()
    instruction_run = instruction_para.add_run("编辑说明：请直接编辑下方内容，但请保持标记符号（===...===）不变，这些标记用于系统识别各个字段。")
    _setup_docx_run(instruction_run)
    _setup_docx_paragraph(instruction_para)
    
    # 添加各个字段的标记和内容
    markers_and_content = [
        ("===TITLE_START===", raw_data.get('title', ''), "===TITLE_END==="),
        ("===CONTENT_TITLE_START===", raw_data.get('content_title', ''), "===CONTENT_TITLE_END==="),
        ("===COVER_SUBTITLE_START===", raw_data.get('cover_subtitle', ''), "===COVER_SUBTITLE_END==="),
        ("===GOLDEN_QUOTE_START===", raw_data.get('golden_quote', ''), "===GOLDEN_QUOTE_END==="),
        ("===CONTENT_START===", raw_data.get('content', ''), "===CONTENT_END===")
    ]
    
    for start_marker, content, end_marker in markers_and_content:
        document.add_paragraph()  # 空行分隔
        
        # 开始标记
        start_para = document.add_paragraph(start_marker)
        _setup_docx_paragraph(start_para)
        
        # 内容
        content_para = document.add_paragraph(content)
        _setup_docx_paragraph(content_para)
        
        # 结束标记
        end_para = document.add_paragraph(end_marker)
        _setup_docx_paragraph(end_para)

    # 统一设置所有段落的字体
    for para in document.paragraphs:
        for run in para.runs:
            _setup_docx_run(run)

    ensure_directory_exists(os.path.dirname(docx_path))
    document.save(docx_path)
    logger.info(f"原始编辑版DOCX已保存: {docx_path}")
    return docx_path


def parse_raw_from_docx(docx_path: str) -> Dict[str, Any]:
    """
    从带标记的DOCX文档解析回原始数据格式
    
    Args:
        docx_path: 包含标记的docx文件路径
        
    Returns:
        Dict[str, Any]: 解析得到的原始数据，包含title、golden_quote、content
        
    Raises:
        FileProcessingError: 解析失败时抛出
    """
    from docx import Document
    
    if not os.path.exists(docx_path):
        raise FileProcessingError(f"DOCX文件不存在: {docx_path}")
    
    try:
        document = Document(docx_path)
        
        # 提取所有段落文本
        paragraphs = []
        for para in document.paragraphs:
            text = para.text.strip()
            if text:  # 只保留非空段落
                paragraphs.append(text)
        
        # 查找标记位置
        title_start = title_end = quote_start = quote_end = content_start = content_end = -1
        content_title_start = content_title_end = cover_subtitle_start = cover_subtitle_end = -1
        
        for i, para_text in enumerate(paragraphs):
            if para_text == "===TITLE_START===":
                title_start = i
            elif para_text == "===TITLE_END===":
                title_end = i
            elif para_text == "===CONTENT_TITLE_START===":
                content_title_start = i
            elif para_text == "===CONTENT_TITLE_END===":
                content_title_end = i
            elif para_text == "===COVER_SUBTITLE_START===":
                cover_subtitle_start = i
            elif para_text == "===COVER_SUBTITLE_END===":
                cover_subtitle_end = i
            elif para_text == "===GOLDEN_QUOTE_START===":
                quote_start = i
            elif para_text == "===GOLDEN_QUOTE_END===":
                quote_end = i
            elif para_text == "===CONTENT_START===":
                content_start = i
            elif para_text == "===CONTENT_END===":
                content_end = i
        
        # 验证标记完整性
        if title_start == -1 or title_end == -1:
            raise ValueError("缺少TITLE标记")
        if quote_start == -1 or quote_end == -1:
            raise ValueError("缺少GOLDEN_QUOTE标记")
        if content_start == -1 or content_end == -1:
            raise ValueError("缺少CONTENT标记")
            
        # 提取各字段内容
        title = '\n'.join(paragraphs[title_start + 1:title_end]).strip()
        golden_quote = '\n'.join(paragraphs[quote_start + 1:quote_end]).strip()
        content = '\n'.join(paragraphs[content_start + 1:content_end]).strip()
        content_title = ''
        if content_title_start != -1 and content_title_end != -1 and content_title_end > content_title_start:
            content_title = '\n'.join(paragraphs[content_title_start + 1:content_title_end]).strip()
        cover_subtitle = ''
        if cover_subtitle_start != -1 and cover_subtitle_end != -1 and cover_subtitle_end > cover_subtitle_start:
            cover_subtitle = '\n'.join(paragraphs[cover_subtitle_start + 1:cover_subtitle_end]).strip()
        
        result = {
            'title': title,
            'content_title': content_title,
            'cover_subtitle': cover_subtitle,
            'golden_quote': golden_quote,
            'content': content
        }
        
        logger.info(f"从DOCX解析原始数据成功: {docx_path}")
        return result
        
    except Exception as e:
        logger.error(f"解析DOCX文件失败: {str(e)}")
        raise FileProcessingError(f"解析DOCX文件失败: {str(e)}")


def _setup_docx_style(document, style_name: str = 'Normal'):
    """设置DOCX文档样式"""
    try:
        from docx.enum.text import WD_LINE_SPACING
        from docx.oxml.ns import qn
        
        style = document.styles[style_name]
        style.font.name = '宋体'
        
        # 设置东亚字体
        if hasattr(style, 'element') and style.element is not None:
            rPr = style.element.rPr
            if rPr is not None and rPr.rFonts is not None:
                rPr.rFonts.set(qn('w:eastAsia'), '宋体')
        
        # 行距 1.5 倍
        if style.paragraph_format is not None:
            try:
                style.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
            except Exception:
                style.paragraph_format.line_spacing = 1.5
    except Exception:
        pass  # 样式设置失败也不影响功能


def _setup_docx_run(run):
    """设置DOCX文本run的字体"""
    try:
        from docx.oxml.ns import qn
        run.font.name = '宋体'
        run._element.rPr.rFonts.set(qn('w:eastAsia'), '宋体')
    except Exception:
        pass


def _setup_docx_paragraph(para, alignment=None):
    """设置DOCX段落的格式"""
    try:
        from docx.enum.text import WD_LINE_SPACING
        if alignment:
            para.alignment = alignment
        para.paragraph_format.line_spacing_rule = WD_LINE_SPACING.ONE_POINT_FIVE
    except Exception:
        try:
            para.paragraph_format.line_spacing = 1.5
        except Exception:
            pass